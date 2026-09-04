"""
Blood test report extraction
============================
A separate module from the payment side. Both share the OCR layer in core.py
and nothing else -- a lab report is a different shape of document, so it gets
its own rules, its own result type and its own exports.

    python medical.py     run the end-to-end self test

Why it is not just more FieldRules
----------------------------------
A payment receipt is label-and-value: "UTR No: 3128...". A lab report is a
*table*: every row is `name, result, unit, reference range`, and the ranges are
printed on the page. So the parser here works row-wise and, crucially, prefers
the reference range the lab printed over any built-in one -- ranges vary by
laboratory, method, age and sex, and the report's own range is the only one
that is definitely right for that sample.

Scope
-----
This extracts and transcribes. A value is flagged only by comparing it against
the range printed beside it, which is arithmetic, not interpretation. Nothing
here diagnoses anything or suggests what to do about a result.
"""

from __future__ import annotations

import csv
import io
import json
import re
import sys
from collections import Counter
from dataclasses import dataclass, field
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Sequence

import core
from core import FieldRule, Line, Match, OcrResult, _clean_plain

REPORT_TITLE = "Blood Test Report Extract"

DISCLAIMER = (
    "Transcribed from the uploaded report. A value is marked high or low only "
    "by comparing it with the reference range printed on that same report -- "
    "this is not a medical opinion. Check anything that matters against the "
    "original document and with the doctor who ordered the test."
)


# =========================================================================== #
#  1. REPORT METADATA
# =========================================================================== #

def _clean_name(value: str) -> str:
    """
    Tidy a patient name read off a crowded header line.

    Labs pack several fields onto one row -- "NAME : MRS.SUBBU LAKSHMI Age:60
    Yr/F" -- so the value pattern happily runs on into the next label. Cutting
    at the first following label word keeps the name to itself.
    """
    text = _clean_plain(value)
    # A value that itself STARTS with a label word ("Age / Sex", "Date of
    # Report") is not a name at all -- it is the *next* column's header
    # bleeding in, on a report that prints every column's header on one
    # shared row ("Patient's Name  Age / Sex") and every column's value on
    # the row below. Matching "Patient's Name" here correctly finds that
    # header row, and the rest of that same row ("Age / Sex") satisfies the
    # value pattern well enough to be accepted as a name unless rejected
    # here -- an empty return fails core.py's post-normalise check, which
    # lets the ordinary next-line lookahead reach the real name one row
    # down instead of settling for "Age /". Confirmed on a real report.
    # "of"/"the" catch the same family of false positive one step further
    # removed: the bare "name" alias (needed for reports with no fuller
    # label at all) matches just as readily inside "Name of the Lab" or
    # "Name of Doctor" -- an unrelated field, not the patient's -- and after
    # "Name" is consumed as the label, what is left starts with "of the ...".
    # A real name never starts with either word. Confirmed on a real report:
    # this specific phrase won out over the correct name at equal
    # confidence, purely because "name" is not anchored to any particular
    # field the way a fuller label like "patient name" is.
    # "No" (as in "Sl.No." / a serial-number column) catches a third case in
    # the same family -- a name-shaped fragment that is really a stray table
    # header the forward lookahead reached before ever finding the real
    # name. Rejecting it here is also what lets a rarer label-below-value
    # layout (see FieldRule.lookbehind) get a chance to look upward instead,
    # since that only triggers once nothing usable turned up below the label
    # at all.
    if re.match(r"^(?:Age|Sex|Gender|Ref|Regist\w*|Reg|Lab|Date|UHID|Bill|"
                r"Sample|Specimen|Patient|Report|of|the|No|Collected|"
                r"Physician|Referral|Doctor|Consultant|Branch|Mobile|Email|"
                r"Phone)\b", text, re.IGNORECASE):
        return ""
    # Regist\w*, not Reg\b: "Reg" only matches a complete word, so it missed
    # "Registered Date" entirely (no boundary between "g" and the "i" that
    # follows) and let a name run on into that neighboring column. Confirmed
    # on a real report.
    text = re.split(r"\s+(?=(?:Age|Sex|Gender|Ref|Regist\w*|Reg|Lab|Date|"
                    r"UHID|Bill|Sample|Specimen|Patient|Report|Dt|Collected|"
                    r"Physician|Referral|Doctor|Consultant|Branch|Mobile|"
                    r"Email|Phone)\b)", text,
                    maxsplit=1, flags=re.IGNORECASE)[0]
    # A leading "Patient :" survives only on values the scan_patterns below
    # produced -- their own match has to include the label text, since there
    # is no separate age/sex label on that layout to anchor a normal match
    # against and bare "patient" cannot be added as an alias (see the comment
    # on patient_name). Stripped here rather than there because normalise is
    # the one step every match, labelled or scanned, always passes through.
    text = re.sub(r"^patient\s*[:\-]\s*", "", text, flags=re.IGNORECASE)
    text = text.strip(" :.-,")
    # A genuine name is essentially never this short once OCR misreads and
    # honorifics are accounted for -- but an isolated fragment the loose
    # next-line/prev-line scan picked up off some unrelated nearby text
    # often is: a serial-number column header, a month abbreviation out of a
    # date ("04-Apr-26"). Rather than block-list every such fragment one at
    # a time as each turns up on a new report, reject anything this short
    # outright and let scanning continue past it. Confirmed on a real
    # report where this was the difference between two separate false
    # positives and the real name three lines away.
    if len(text) < 5:
        return ""
    return text


META_RULES: list[FieldRule] = [
    FieldRule(
        key="patient_name", label="Patient Name",
        # No bare "patient" or "name": the label regex lets an optional "id"
        # follow an alias, so "patient" happily matches the box "Patient ID" and
        # then walks off to grab whatever text sits further along that row.
        # "name" last: aliases are tried longest-first, and it is anchored to the
        # start of a box, so the "TEST NAME" column header cannot match it.
        # "pt's" is the routine abbreviation on many Indian lab reports; OCR
        # just as routinely drops or misreads its apostrophe, so all three
        # spellings are listed rather than relying on the loose fallback to
        # bridge them. Confirmed missing entirely (no match at all, not even
        # a weak one) on a real report printing "Pt's NAME".
        aliases=["patient name", "patients name", "patient's name",
                 "name of patient", "pt's name", "pts name", "pt s name",
                 "name"],
        # /  covers "W/O", "S/O", "D/O", "C/O" -- Wife/Son/Daughter/Care Of,
        # printed as part of the name itself on many Indian reports ("MRS.
        # CHANDRA W/O ANANDAN"). Without it the match simply stops dead at
        # the slash, truncating a real name to its first word. Confirmed on
        # a real report.
        value_pattern=r"(?:M/?R?S?\.?\s*)?[A-Za-z][A-Za-z./\s]{2,50}",
        # Not strict: several labs print "Patient Name  MR. X" inside one box,
        # and a strict label would refuse to look at the rest of that box. The
        # aliases are long and distinctive enough to be safe unanchored.
        strict_label=False, multi_segment_value=True, normalise=_clean_name,
        # Some reports print "Name" as its own row underneath the actual
        # patient-info line instead of above it. Scoped to this field only --
        # see FieldRule.lookbehind.
        lookbehind=2,
        # Only reached when every alias above found nothing -- covers a
        # layout with no "Name" label at all, just a bare "Patient :" that
        # cannot be a normal alias (see above) followed directly by the name
        # and an "(age/sex)" parenthetical with no label of its own either
        # ("Patient : Mrs.Meenakumari (60/F)"). Requiring that exact trailing
        # shape is what makes "patient" safe to match here where it is not as
        # a plain alias: a genuine "Patient ID : AJH2156" row has no "(nn/M)"
        # or "(nn/F)" anywhere after it, so it never satisfies this pattern.
        # Confirmed on a real report.
        scan_patterns=[
            (r"(?i)patient\s*:\s*[A-Za-z][A-Za-z./\s]{2,40}?"
             r"(?=\s*\(\d{1,3}\s*/\s*(?:M|F|Male|Female)\s*\))", 1.0),
        ],
    ),
    FieldRule(
        key="patient_id", label="Patient / UHID",
        aliases=["patient id", "uhid", "mrn", "reg no", "registration no",
                 "lab id", "lab no", "sample id", "accession no", "barcode"],
        value_pattern=r"[A-Za-z0-9][A-Za-z0-9\-/]{3,24}",
        normalise=lambda v: re.sub(r"\s", "", v).upper(), multi=True,
    ),
    FieldRule(
        key="age_sex", label="Age / Sex",
        aliases=["age/sex", "age / sex", "age & sex", "age", "sex", "gender"],
        # The separator also accepts a lone "1" -- a real, repeated OCR
        # misread of the printed "/" in "55 / Female" -- but only when M/F
        # actually follows, so a genuine "1" elsewhere in the row is never
        # swallowed. multi_segment_value matters here specifically because
        # OCR often boxes the age, the separator and the sex as three
        # separate segments on one visual row; without it, "55" alone (a
        # complete match on its own, since sex is optional) wins and the row
        # is never rejoined with the segment carrying the sex. Confirmed on a
        # real report: age_sex came back as "55" with the sex silently lost,
        # which then picked the wrong (Male) half of a sex-split reference
        # range for an unrelated Haemoglobin row.
        value_pattern=r"\d{1,3}\s*(?:y(?:rs?|ears?)?)?\s*(?:[/,]|1(?=\s*[MF]))?\s*(?:M|F|Male|Female)?",
        multi_segment_value=True,
        normalise=_clean_plain,
        # Same layout as patient_name's scan_patterns above and reached for
        # the same reason: a report that gives age/sex only as "(60/F)" right
        # after the name has no "Age"/"Sex" word anywhere for the aliases
        # above to find. The parentheses are matched via look-around so they
        # never end up as part of the captured value. Confirmed on a real
        # report.
        scan_patterns=[
            (r"(?i)(?<=\()\d{1,3}\s*/\s*(?:M|F|Male|Female)(?=\))", 1.0),
            # A report that gives age/sex as its own bare line -- "70/FEMALE"
            # on a row by itself, no "Age"/"Sex" word and no parentheses
            # either -- has nothing at all for the parenthetical pattern
            # above, or the aliases, to anchor to. Anchored to the WHOLE
            # line (^...$, not just found somewhere in it) specifically so
            # this cannot misfire on some unrelated "70/110"-shaped number
            # pair sitting inside a longer line -- a genuine age/sex report
            # line really is just this and nothing else. Confirmed on a
            # real report.
            (r"(?i)^\s*\d{1,3}\s*/\s*(?:M|F|Male|Female)\s*$", 1.0),
        ],
    ),
    FieldRule(
        key="referred_by", label="Referred By",
        aliases=["referred by", "ref. doctor", "ref doctor", "ref.by dr",
                 "ref by dr", "ref. by", "ref by", "referring doctor",
                 "consultant", "doctor"],
        value_pattern=r"[A-Za-z][A-Za-z.\s]{2,50}",
        strict_label=False, multi_segment_value=True, normalise=_clean_plain,
    ),
    FieldRule(
        key="lab_name", label="Laboratory",
        aliases=["laboratory", "lab name", "diagnostics", "pathology"],
        value_pattern=r"[A-Za-z][A-Za-z0-9.,&\s]{3,60}",
        normalise=_clean_plain,
    ),
    FieldRule(
        key="collected_on", label="Sample Collected",
        aliases=["collected on", "collection date", "sample collected",
                 "sample date", "drawn on"],
        value_pattern=core.DATE_PATTERN, normalise=_clean_plain,
    ),
    FieldRule(
        key="reported_on", label="Reported On",
        aliases=["reported on", "report date", "rpt date", "reporting date",
                 "released on", "date of report"],
        value_pattern=core.DATE_PATTERN, normalise=_clean_plain,
    ),
    FieldRule(
        key="sample_type", label="Sample Type",
        aliases=["sample type", "specimen", "sample"],
        value_pattern=r"[A-Za-z][A-Za-z\s]{2,30}", normalise=_clean_plain,
    ),
]

META_ORDER = [r.key for r in META_RULES]


# =========================================================================== #
#  2. ANALYTES
# =========================================================================== #

@dataclass
class Analyte:
    """One measurable quantity and the fallback range for an adult."""
    key: str
    name: str
    aliases: Sequence[str]
    unit: str = ""
    low: float | None = None
    high: float | None = None
    panel: str = "Other"
    note: str = ""
    # Female range, where it genuinely differs. Used only when the report prints
    # no range of its own AND the patient's sex was read off the header.
    low_f: float | None = None
    high_f: float | None = None

    def range_for(self, sex: str | None) -> tuple[float | None, float | None]:
        if sex == "F" and (self.low_f is not None or self.high_f is not None):
            return self.low_f, self.high_f
        return self.low, self.high


def detect_sex(meta: dict[str, str]) -> str | None:
    """
    Read the patient's sex out of whatever the header printed.

    Handles "42 Y / F", "61 / MALE", "60 Yr/F", "Male". Returns None when it is
    absent or ambiguous -- guessing here would produce a confident verdict
    against the wrong range, which is worse than declining to judge.
    """
    blob = " ".join(v for k, v in meta.items()
                    if k in ("age_sex", "patient_name") and v)
    if not blob:
        return None
    if re.search(r"\bfemale\b|\bF\b|/\s*F\b", blob, re.IGNORECASE):
        if not re.search(r"\bmale\b", blob, re.IGNORECASE):
            return "F"
    if re.search(r"\bmale\b|\bM\b|/\s*M\b", blob, re.IGNORECASE):
        return "M"
    return None


# Fallback ranges only -- typical adult values, used when the report does not
# print its own. Several genuinely differ by sex or method, which is exactly why
# the printed range always wins.
ANALYTES: list[Analyte] = [
    # ---- Complete blood count -------------------------------------------- #
    Analyte("haemoglobin", "Haemoglobin",
            ["haemoglobin", "hemoglobin", "hb", "hgb"], "g/dL", 13.0, 17.0,
            "Complete Blood Count", "range differs by sex",
            low_f=12.0, high_f=15.0),
    Analyte("rbc", "RBC Count", ["rbc count", "rbc", "red blood cell",
            "total rbc"], "million/µL", 4.5, 5.5, "Complete Blood Count",
            "range differs by sex", low_f=3.8, high_f=4.8),
    Analyte("wbc", "WBC / Total Leucocyte Count",
            ["total leucocyte count", "total leukocyte count", "wbc count",
             "tlc", "wbc", "leucocyte count"], "/µL", 4000, 11000,
            "Complete Blood Count"),
    Analyte("platelet", "Platelet Count",
            ["platelet count", "platelets", "plt"], "/µL", 150000, 450000,
            "Complete Blood Count"),
    Analyte("hct", "Haematocrit / PCV",
            ["haematocrit", "hematocrit", "pcv", "packed cell volume", "hct"],
            "%", 40.0, 50.0, "Complete Blood Count", "range differs by sex",
            low_f=36.0, high_f=46.0),
    Analyte("mcv", "MCV", ["mcv", "mean corpuscular volume"], "fL", 83.0, 101.0,
            "Complete Blood Count"),
    Analyte("mch", "MCH", ["mch", "mean corpuscular haemoglobin"], "pg", 27.0, 32.0,
            "Complete Blood Count"),
    Analyte("mchc", "MCHC", ["mchc"], "g/dL", 31.5, 34.5, "Complete Blood Count"),
    # Split into two entries, not one -- RDW-CV (%) and RDW-SD (fL) are
    # printed as separate rows measuring different things. A single "rdw"
    # entry with only "rdw-cv" as an extra alias meant "RDW-SD" matched via
    # the bare "rdw" substring anyway (it contains "rdw"), so the SD reading
    # landed as a second, wrongly-unitted "RDW" value instead of its own
    # row. The bare "rdw" alias stays only on -CV, since that is what an
    # unqualified "RDW" on a report means. Confirmed on a real report.
    Analyte("rdw_cv", "RDW-CV", ["rdw-cv", "rdw cv", "rdw(cv)", "rdw"], "%",
            11.6, 14.0, "Complete Blood Count"),
    Analyte("rdw_sd", "RDW-SD", ["rdw-sd", "rdw sd", "rdw(sd)"], "fL",
            39.0, 46.0, "Complete Blood Count"),
    Analyte("neutrophils", "Neutrophils", ["neutrophils", "neutrophil"], "%",
            40.0, 80.0, "Differential Count"),
    Analyte("lymphocytes", "Lymphocytes", ["lymphocytes", "lymphocyte"], "%",
            20.0, 40.0, "Differential Count"),
    Analyte("eosinophils", "Eosinophils", ["eosinophils", "eosinophil"], "%",
            1.0, 6.0, "Differential Count"),
    Analyte("monocytes", "Monocytes", ["monocytes", "monocyte"], "%",
            2.0, 10.0, "Differential Count"),
    Analyte("basophils", "Basophils", ["basophils", "basophil"], "%",
            0.0, 2.0, "Differential Count"),
    Analyte("esr", "ESR", ["esr", "erythrocyte sedimentation rate"], "mm/hr",
            0.0, 20.0, "Complete Blood Count"),

    # ---- Diabetes --------------------------------------------------------- #
    # The alias lists carry every spelling seen on real reports, including the
    # labs' own misspellings ("cholestrol", "triglericids"). A typo that is
    # printed on thousands of reports is not worth being pedantic about.
    Analyte("glucose_f", "Glucose (Fasting)",
            ["blood sugar (fasting)", "blood sugar(fasting)", "fasting blood sugar",
             "blood sugar fasting", "glucose fasting", "fasting glucose",
             "bl.sugar (f)", "bl.sugar(f)", "bl sugar (f)", "b.sugar (f)",
             "sugar (f)", "fbs", "sugar fasting",
             "glucose (f)", "glucose(f)", "glucose - f", "glucose-f",
             "glucose ( f)", "plasma sugar (f)", "plasma sugar(f)",
             "fasting plasma glucose"], "mg/dL",
            70.0, 100.0, "Diabetes"),
    Analyte("glucose_pp", "Glucose (Post Prandial)",
            ["post prandial blood sugar", "blood sugar (p.p)", "blood sugar(p.p)",
             "blood sugar (pp)", "blood sugar(pp)", "glucose post prandial",
             "bl.sugar (pp)", "bl.sugar(pp)", "bl sugar (pp)", "b.sugar(pp)",
             "sugar (pp)", "pp glucose", "ppbs", "blood sugar pp",
             "post prandial", "postprandial", "blood sugar postprandial",
             "glucose postprandial",
             "glucose (pp)", "glucose(pp)",
             "glucose - pp", "glucose-pp", "glucose ( pp)",
             "plasma sugar (pp", "plasma sugar(pp", "before lunch plasma glucose"],
            "mg/dL", 70.0, 140.0, "Diabetes"),
    Analyte("glucose_r", "Glucose (Random)",
            ["random blood sugar", "blood sugar (random)", "glucose random",
             "rbs", "random blood glucose", "glucose (random)",
             "glucose - random"], "mg/dL", 70.0, 140.0, "Diabetes"),
    Analyte("hba1c", "HbA1c",
            ["hba1c (biorad)", "hba1c(biorad)", "hba1c (bio-rad)", "hba1c",
             "hb a1c", "hba1 c", "hba1_c", "glycated haemoglobin",
             "glycosylated haemoglobin",
             # Bare, not just the full "glycosylated haemoglobin" phrase --
             # one report split the test's own name across two lines
             # ("GLYCOSYLATED" / <value + a tier annotation> / "HAEMOGLOBIN
             # - HbA1c"), so "haemoglobin" was never on the same line as
             # "glycosylated" at all. Confirmed on a real report where this
             # was the difference between the result being silently
             # dropped -- the row also contained "NON DIABETIC:", which
             # made it look like pure interpretation-tier noise once it
             # went entirely unrecognised -- and being read correctly.
             "glycosylated",

             # Spelled out in full rather than abbreviated -- "Hemoglobin
             # A1c 6.16 %" / "Hemoglobin A1 c : 5.90 %". Without these, the
             # bare "haemoglobin"/"hemoglobin" alias below (a different
             # test -- plain CBC Hemoglobin) matched first and the row came
             # back as a CBC Hemoglobin reading: wrong unit (g/dL vs %),
             # wrong range (13-17 vs ~4-6), and since HbA1c and Hemoglobin
             # move in opposite clinical directions, sometimes an inverted
             # flag too (a high HbA1c reading as a low Hemoglobin).
             # Confirmed on two different real reports.
             "hemoglobin a1c", "hemoglobin a1 c", "hemoglobin a1",
             "haemoglobin a1c", "haemoglobin a1 c", "haemoglobin a1"],
            "%", 4.0, 5.6, "Diabetes"),
    # Derived from HbA1c, printed under many names. No built-in range: labs
    # publish it against their own bands, so a value with no printed range is
    # reported without a verdict rather than judged against a guess.
    Analyte("eag", "Estimated Average Glucose",
            ["estimated average glucose", "estimated glucose level",
             "mean blood glucose", "average blood glucose",
             "means glucose value", "mean glucose value", "means glucose",
             "mean glucose", "mean plasma glucose",
             "estimation of mean blood glucose", "estimated average blood glucose",
             "estimated avg glucose", "eag", "eab", "abg"],
            "mg/dL", None, None, "Diabetes"),

    # ---- Lipids ----------------------------------------------------------- #
    Analyte("chol_total", "Total Cholesterol",
            ["serum cholesterol", "serum cholestrol", "total cholesterol",
             "total cholestrol", "cholesterol total", "cholesterol",
             "cholestrol"], "mg/dL", None, 200.0, "Lipid Profile"),
    # "Cholesterol, HDL"/"Cholesterol, LDL"/"Cholesterol, VLDL" -- the
    # qualifier printed AFTER the word, comma-separated -- are added
    # explicitly rather than relying on the fuzzy fallback to bridge them.
    # The comma breaks an exact match against "hdl cholesterol" (space, no
    # comma, different word order), so that fails, and the loop moves on to
    # try shorter aliases -- where bare "cholesterol" (a real alias of Total
    # Cholesterol) matches "Cholesterol," as a plain substring and wins
    # first, before the fuzzy fallback that WOULD have handled the comma
    # ever gets a chance to run. Confirmed on a real report: three distinct
    # results (HDL/LDL/VLDL) all showed up mislabelled as extra "readings"
    # of Total Cholesterol.
    Analyte("hdl", "HDL Cholesterol",
            ["serum hdl cholesterol", "serum hdl", "hdl cholesterol",
             "cholesterol, hdl", "cholesterol,hdl", "cholesterol hdl",
             "cholesterol - hdl", "hdl"],
            "mg/dL", 40.0, None, "Lipid Profile", "range differs by sex",
            low_f=50.0),
    Analyte("ldl", "LDL Cholesterol",
            ["serum ldl cholesterol", "serum ldl", "ldl cholesterol",
             "cholesterol, ldl", "cholesterol,ldl", "cholesterol ldl",
             "cholesterol - ldl", "ldl"],
            "mg/dL", None, 100.0, "Lipid Profile"),
    Analyte("vldl", "VLDL Cholesterol",
            ["serum vldl cholesterol", "serum vldl", "vldl cholesterol",
             "cholesterol, vldl", "cholesterol,vldl", "cholesterol vldl",
             "cholesterol - vldl", "vldl"],
            "mg/dL", None, 30.0, "Lipid Profile"),
    Analyte("triglycerides", "Triglycerides",
            ["serum triglycerides", "serum triglericids", "serum triglycerids",
             "triglycerides", "triglycerids", "triglericids", "triglyceride",
             "tg"], "mg/dL", None, 150.0, "Lipid Profile"),
    # A distinct entry rather than leaving this to fall through to Total
    # Cholesterol: "Total Cholesterol/HDL" is a ratio, unitless, with its own
    # row on the report. Without this the row's alias ("total cholesterol")
    # matched as a PREFIX of "total cholesterol/hdl" -- "/" was accepted as a
    # word boundary the same as a space -- so the ratio silently overwrote the
    # real cholesterol reading under one shared key. Confirmed on a real report.
    Analyte("chol_hdl_ratio", "Total Cholesterol / HDL Ratio",
            ["total cholesterol/hdl", "total cholesterol / hdl",
             "cholesterol/hdl ratio", "tc/hdl ratio", "chol/hdl ratio",
             # Reversed word order -- "Cholesterol Total/Cholesterol HDL",
             # not "Total Cholesterol/HDL" -- from the same report that
             # needed HDL/LDL/VLDL's own space/dash-separator variants.
             # Without it this matched bare "cholesterol" and overwrote the
             # real Total Cholesterol reading, same failure mode as the
             # original bug this analyte was added to fix. Confirmed on a
             # real report.
             "cholesterol total/cholesterol hdl",
             "cholesterol total / cholesterol hdl",
             "coronary risk ratio-1", "coronary risk ratio -1",
             "coronary risk ratio - 1", "coronary risk ratio 1",
             "coronary risk ratio-i"],
            "", None, None, "Lipid Profile"),
    # Same family as chol_hdl_ratio above but a different formula (LDL/HDL, not
    # T.Chol/HDL) -- printed under its own row on reports that show both.
    # Previously absent entirely, so its alias "hdl" matched inside the row and
    # silently overwrote the real HDL Cholesterol reading. Confirmed on a real
    # report.
    Analyte("ldl_hdl_ratio", "LDL / HDL Ratio",
            ["ldl/hdl ratio", "ldl / hdl ratio", "ldl/hdl", "ldl/ hdl",
             "ldl/hdl cholesterol ratio", "ldl / hdl cholesterol ratio",
             # Same reversed-order variant as chol_hdl_ratio above --
             # "Cholesterol LDL/Cholesterol HDL", not "LDL/HDL Cholesterol
             # Ratio". Confirmed on a real report.
             "cholesterol ldl/cholesterol hdl",
             "cholesterol ldl / cholesterol hdl",
             "coronary risk ratio-ii", "coronary risk ratio -ii",
             "coronary risk ratio - ii", "coronary risk ratio ii",
             "coronary risk ratio-2", "coronary risk ratio -2",
             "coronary risk ratio - 2", "coronary risk ratio 2"],
            "", None, None, "Lipid Profile"),
    # HDL/LDL is the inverse of the ratio above and prints as its own row on
    # reports that show both directions. Without its own entry, "HDL / LDL
    # Cholesterol Ratio" starts with "HDL" and the hdl_chol analyte's own
    # alias silently swallows it as a second, wrong HDL Cholesterol reading
    # -- the same failure mode ldl_hdl_ratio was added to fix, just the other
    # way round. Confirmed on a real report.
    Analyte("hdl_ldl_ratio", "HDL / LDL Ratio",
            ["hdl/ldl ratio", "hdl / ldl ratio", "hdl/ldl", "hdl/ ldl",
             "hdl/ldl cholesterol ratio", "hdl / ldl cholesterol ratio"],
            "", None, None, "Lipid Profile"),
    # Total Cholesterol minus HDL -- a distinct printed row, not a stand-in for
    # Total Cholesterol. Previously absent, so its alias "cholesterol" matched
    # inside "Non HDL Cholesterol" and overwrote the real Total Cholesterol
    # reading. Confirmed on a real report. 130 mg/dL is the commonly printed
    # desirable ceiling, used only when a report gives no range of its own.
    Analyte("non_hdl_chol", "Non-HDL Cholesterol",
            ["non hdl cholesterol", "non-hdl cholesterol", "non hdl cholestrol",
             "non-hdl cholestrol"],
            "mg/dL", None, 130.0, "Lipid Profile"),

    # ---- Liver ------------------------------------------------------------ #
    Analyte("bilirubin_t", "Bilirubin (Total)",
            ["total bilirubin", "bilirubin total", "bilirubin"], "mg/dL",
            0.3, 1.2, "Liver Function"),
    Analyte("bilirubin_d", "Bilirubin (Direct)",
            ["direct bilirubin", "bilirubin direct", "bilirubin-direct",
             "conjugated bilirubin"],
            "mg/dL", 0.0, 0.3, "Liver Function"),
    Analyte("bilirubin_i", "Bilirubin (Indirect)",
            ["indirect bilirubin", "bilirubin indirect", "bilirubin-indirect",
             "unconjugated bilirubin"],
            "mg/dL", 0.2, 0.8, "Liver Function"),
    Analyte("sgpt", "SGPT / ALT", ["sgpt", "alt", "alanine aminotransferase",
            "sgpt (alt)"], "U/L", 7.0, 56.0, "Liver Function"),
    Analyte("sgot", "SGOT / AST", ["sgot", "ast", "aspartate aminotransferase",
            "sgot (ast)"], "U/L", 5.0, 40.0, "Liver Function"),
    # A derived row, not a third liver enzyme -- without its own entry,
    # "SGOT/SGPT Ratio" starts with "sgot" and gets swallowed as a second,
    # wrong SGOT/AST reading. Same failure mode as ldl_hdl_ratio above.
    # Confirmed on a real report.
    Analyte("sgot_sgpt_ratio", "SGOT / SGPT Ratio",
            ["sgot/sgpt ratio", "sgot / sgpt ratio", "sgot/sgpt",
             "ast/alt ratio", "ast / alt ratio"],
            "", None, None, "Liver Function"),
    Analyte("alp", "Alkaline Phosphatase", ["alkaline phosphatase", "alp"],
            "U/L", 44.0, 147.0, "Liver Function"),
    Analyte("protein_total", "Total Protein", ["total protein", "protein total"],
            "g/dL", 6.4, 8.3, "Liver Function"),
    Analyte("albumin", "Albumin", ["albumin"], "g/dL", 3.5, 5.2, "Liver Function"),
    Analyte("globulin", "Globulin", ["globulin"], "g/dL", 2.0, 3.5,
            "Liver Function"),
    # Same failure mode again -- "Albumin/Globulin Ratio" starts with
    # "albumin" and would otherwise read as a second Albumin value. Confirmed
    # on a real report.
    Analyte("ag_ratio", "Albumin / Globulin Ratio",
            ["albumin/globulin ratio", "albumin / globulin ratio",
             "albumin/globulin", "a/g ratio", "a: g ratio"],
            "", None, None, "Liver Function"),

    # ---- Kidney and electrolytes ------------------------------------------ #
    Analyte("urea", "Urea", ["blood urea", "urea"], "mg/dL", 15.0, 45.0,
            "Kidney Function"),
    Analyte("bun", "Blood Urea Nitrogen", ["blood urea nitrogen", "bun"],
            "mg/dL", 7.0, 20.0, "Kidney Function"),
    Analyte("creatinine", "Creatinine", ["serum creatinine", "creatinine"],
            "mg/dL", 0.6, 1.3, "Kidney Function"),
    # Two more derived rows with the same failure mode as the liver-panel
    # ratios above -- "Urea / Sr.Creatinine Ratio" starts with "urea" and
    # "Bun/Creatinine Ratio" starts with "bun", so without their own entries
    # each reads as a second, wrong Urea/BUN value. Confirmed on real reports.
    Analyte("urea_creatinine_ratio", "Urea / Creatinine Ratio",
            ["urea/creatinine ratio", "urea / creatinine ratio",
             "urea/sr.creatinine ratio", "urea / sr.creatinine ratio",
             "urea/creatinine"],
            "", None, None, "Kidney Function"),
    Analyte("bun_creatinine_ratio", "BUN / Creatinine Ratio",
            ["bun/creatinine ratio", "bun / creatinine ratio",
             "bun/cr ratio", "bun/creatinine"],
            "", None, None, "Kidney Function"),
    Analyte("uric_acid", "Uric Acid", ["uric acid", "serum uric acid"], "mg/dL",
            3.5, 7.2, "Kidney Function"),
    # Previously unlisted, so every spelling variant of this very common test
    # ("Est. Glomerular Filtration Rate", "Est.Glomerular Filtration Rate",
    # "Est. Glomerular Filtration Rate Serum") landed as three separate
    # "unknown" rows instead of being recognised -- and deduped -- as one
    # test read multiple times. Confirmed on a real report.
    Analyte("egfr", "Estimated GFR",
            ["estimated glomerular filtration rate",
             "est. glomerular filtration rate", "est glomerular filtration rate",
             "estimated gfr", "est. gfr", "est gfr", "egfr"],
            "mL/min/1.73m2", 90.0, None, "Kidney Function"),
    Analyte("sodium", "Sodium", ["sodium", "na+", "serum sodium"], "mEq/L",
            135.0, 145.0, "Electrolytes"),
    Analyte("potassium", "Potassium", ["potassium", "k+", "serum potassium"],
            "mEq/L", 3.5, 5.1, "Electrolytes"),
    Analyte("chloride", "Chloride", ["chloride", "cl-"], "mEq/L", 98.0, 107.0,
            "Electrolytes"),
    Analyte("calcium", "Calcium", ["calcium", "serum calcium"], "mg/dL",
            8.5, 10.5, "Electrolytes"),

    # ---- Thyroid ---------------------------------------------------------- #
    Analyte("tsh", "TSH", ["tsh", "thyroid stimulating hormone"], "µIU/mL",
            0.4, 4.0, "Thyroid Profile"),
    Analyte("t3", "T3 (Total)", ["total t3", "t3 total", "t3",
            "triiodothyronine"], "ng/dL", 80.0, 200.0, "Thyroid Profile"),
    Analyte("t4", "T4 (Total)", ["total t4", "t4 total", "t4", "thyroxine"],
            "µg/dL", 5.1, 14.1, "Thyroid Profile"),
    Analyte("ft3", "Free T3", ["free t3", "ft3"], "pg/mL", 2.3, 4.2,
            "Thyroid Profile"),
    Analyte("ft4", "Free T4", ["free t4", "ft4"], "ng/dL", 0.8, 1.8,
            "Thyroid Profile"),

    # ---- Vitamins and iron ------------------------------------------------ #
    Analyte("vit_d", "Vitamin D (25-OH)",
            ["vitamin d", "25 hydroxy vitamin d", "25-oh vitamin d",
             "vitamin d3", "25(oh)d"], "ng/mL", 30.0, 100.0, "Vitamins"),
    Analyte("vit_b12", "Vitamin B12", ["vitamin b12", "vit b12", "b12",
            "cobalamin"], "pg/mL", 200.0, 900.0, "Vitamins"),
    Analyte("ferritin", "Ferritin", ["ferritin", "serum ferritin"], "ng/mL",
            20.0, 250.0, "Iron Studies"),
    Analyte("iron", "Serum Iron", ["serum iron", "iron"], "µg/dL", 60.0, 170.0,
            "Iron Studies"),
    Analyte("tibc", "TIBC", ["tibc", "total iron binding capacity"], "µg/dL",
            240.0, 450.0, "Iron Studies"),
    Analyte("crp", "CRP", ["c reactive protein", "crp", "hs-crp"], "mg/L",
            0.0, 6.0, "Inflammation"),

    # ---- Tumor markers ----------------------------------------------------- #
    # Unlisted before, so this fell through to the unknown-row path, where
    # the "125" printed as part of the test's OWN name/abbreviation --
    # "Cancer Antigen 125 (CA-125)" -- was mistaken for the result itself
    # (the first number found scanning the row). See the parenthetical-echo
    # handling in _find_analyte/_parse_unknown_row for the actual fix; this
    # entry additionally lets a report using this name get a real range and
    # flag instead of showing up as an unjudged "unknown" row. Confirmed
    # reproducible across two different labs' reports.
    # Short forms ("ca-125", "ca125", "ca 125") are deliberately NOT aliases
    # here. A report's own methodology paragraph routinely opens with prose
    # like "CA 125 is a second generation assay for the detection of...",
    # and that sentence matched this test's name just as readily as the
    # real result row -- confirmed on a real report, reporting a fabricated
    # value read out of the explanation text instead. Digit-folding in the
    # fuzzy fallback (1->i, 2->z, 5->s, for OCR look-alikes) made this worse,
    # not better: even with the exact short alias removed, "ca125" folds to
    # "caizs", and "CA 125 IS a second..." folds to "caizsisasecond...",
    # which still starts with "caizs" -- a coincidental collision between a
    # test name containing real digits and the very next words happening to
    # continue the pattern once folded. "cancer antigen 125" is long and
    # distinctive enough that a report's prose is very unlikely to repeat it
    # verbatim, and it is exactly how the real row on that report was
    # printed, so it alone is enough to catch it.
    Analyte("ca125", "Cancer Antigen 125 (CA-125)",
            ["cancer antigen 125", "cancer antigen-125"],
            "U/mL", None, 35.0, "Tumor Markers"),
]

ANALYTES_BY_KEY = {a.key: a for a in ANALYTES}

# Longest aliases first, so "total cholesterol" is tried before "cholesterol"
# and "free t4" before "t4".
_ALIAS_INDEX: list[tuple[str, Analyte]] = sorted(
    ((alias, a) for a in ANALYTES for alias in a.aliases),
    key=lambda pair: len(pair[0]), reverse=True,
)

# Every spelling of a unit seen on the sample reports. Indian labs print "Mg%",
# "MG%", "mgs/dl" and "mg/dL" for the same thing, so all of them are listed.
# Order matters: "mg%" must precede "%" or the shorter token wins first.
UNIT_TOKENS = [
    "million/µl", "million/ul", "millions/cumm", "million/cumm", "lakhs/cumm", "10\\^3/µl",
    "x10\\^3/µl", "µiu/ml", "uiu/ml", "miu/l",
    "mgs/dl", "mgs/dl", "mg/dl", "gms/dl", "g/dl", "ng/dl",
    "µg/dl", "ug/dl", "ng/ml", "pg/ml", "µg/l", "mg/l", "meq/l", "mmol/l",
    "µmol/l", "iu/l", "u/l", "mm/hr", "/cumm", "cumm", "/µl", "/ul", "fl",
    "mg%", "gm%", "gms%", "g%", "pg", "%",
]
# Preceding digit is allowed on purpose -- "115mg/dl" with no space between
# the value and its unit is the common case on a printed or OCR'd report, not
# the exception. Excluding it here (as an excluded-letter-only lookbehind
# would not) used to leave "mg/dl" glued to the number, and that glued letter
# is exactly what made the number regex below back off from "115" to "11":
# its own trailing (?![A-Za-z]) failed against the "m", so it retried with
# one fewer digit, landed on "5" (a digit, satisfying the guard), and
# returned "11" as the result. Confirmed on a real report.
_UNIT_RE = re.compile(r"(?<![A-Za-z])(" + "|".join(UNIT_TOKENS) + r")(?![A-Za-z])",
                      re.IGNORECASE)

# "13.0 - 17.0", "13.0 to 17.0", "< 200", ">= 40", "0.3-1.2"
_RANGE_BETWEEN = re.compile(
    r"(\d+(?:[.,]\d+)?)\s*(?:-|–|—|to)\s*(\d+(?:[.,]\d+)?)", re.IGNORECASE)
_RANGE_BOUND = re.compile(r"(<=?|>=?|upto|up to|less than|greater than)\s*(\d+(?:[.,]\d+)?)",
                          re.IGNORECASE)
# (?<!\S) requires the number to start at a real token boundary -- preceded by
# whitespace or the start of the string. Without it, a method/analyzer name
# like "BIO RAD D-10 Analyzer" reads as containing the number -10 (from
# "D-10"), and because it sits before the real result in the line, a bare
# .search() finds that first and reports the analyzer's model number as the
# patient's result. Confirmed on a real report: this genuinely turned an
# HbA1c of 8.1 into -10 before the fix.
#
# (?![A-Za-z]) is the same guard on the other side. Without it, "Us TSH - 3rd
# Generation 1.58 mIu/ml" reads as containing the number 3 (from "3rd") --
# the digits of an ordinal are a perfectly valid-looking token on their left
# side, and only checking what came before missed that "rd" glued onto the
# right side means it is not a standalone number at all. Confirmed on a real
# report: this genuinely turned a TSH of 1.58 into 3.
#
# The boundary set also includes a colon (regular ":" or full-width "："),
# unlike the plain whitespace-only version above -- "TRIGLYCERIDES(TGL)
# :98.4" and "...：98.4" are both common, unambiguous "label, then value with
# no space" layouts, not a token a real value could be glued onto by
# accident the way a hyphen (still excluded, see D-10 above) can be.
# Confirmed on a real report: a value directly after a bare colon was
# invisible to this regex, so the row silently failed to parse at all even
# though the correct digits were sitting right there.
_NUMBER = re.compile(
    r"(?<![^\s:：])[-+]?\d{1,3}(?:,\d{3})+(?:\.\d+)?(?![A-Za-z])"
    r"|(?<![^\s:：])[-+]?\d+(?:\.\d+)?(?![A-Za-z])")


def _to_float(text: str) -> float | None:
    try:
        return float(text.replace(",", ""))
    except (TypeError, ValueError):
        return None


# =========================================================================== #
#  3. RESULTS
# =========================================================================== #

@dataclass
class AnalyteResult:
    key: str
    name: str
    panel: str
    value: str                       # exactly as read
    numeric: float | None = None
    unit: str = ""
    ref_text: str = ""               # the range as printed
    ref_low: float | None = None
    ref_high: float | None = None
    ref_source: str = "report"       # "report" or "builtin"
    flag: str = "unknown"            # normal | high | low | unknown
    confidence: float = 0.0
    context: str = ""
    note: str = ""


@dataclass
class BloodReport:
    filename: str = ""
    meta: dict[str, str] = field(default_factory=dict)
    results: list[AnalyteResult] = field(default_factory=list)
    full_text: str = ""
    full_text_display: str = ""
    engine: str = ""
    ocr_confidence: float = 0.0
    elapsed: float = 0.0

    @property
    def out_of_range(self) -> list[AnalyteResult]:
        # "check" belongs here too: a value flagged implausible is at least as
        # worth a person's attention as a plain high or low, arguably more so.
        return [r for r in self.results if r.flag in ("high", "low", "check")]

    @property
    def panels(self) -> list[str]:
        seen: list[str] = []
        for r in self.results:
            if r.panel not in seen:
                seen.append(r.panel)
        return seen


# =========================================================================== #
#  4. ROW PARSING
# =========================================================================== #

# Test names suffer the same glyph confusions as ids do: "Bl.Sugar" comes back
# as "BI.Sugar" because lowercase l and uppercase I are the same shape. Folding
# those classes together lets a name match despite the misread.
_FOLD = str.maketrans({"l": "i", "1": "i", "|": "i", "0": "o", "5": "s",
                       "8": "b", "2": "z", "6": "g"})


def _loose(text: str) -> str:
    """Lowercase, drop punctuation and fold look-alike characters."""
    return re.sub(r"[^a-z0-9]", "", text.lower()).translate(_FOLD)


_LOOSE_INDEX: list[tuple[str, Analyte]] = sorted(
    ((_loose(alias), a) for a in ANALYTES for alias in a.aliases),
    key=lambda pair: len(pair[0]), reverse=True,
)

# Specimen/timing qualifiers a lab commonly prints ahead of a test's own name
# -- "Serum Total Cholesterol", "Random Blood Sugar" -- which must not be
# mistaken for a genuinely different word standing in front of the match
# (see prefix_ok in _find_analyte).
_PREFIX_QUALIFIERS = {
    "serum", "plasma", "blood", "urine", "csf", "s", "sr", "b",
    "fasting", "random", "post", "prandial", "pp", "f", "venous", "capillary",
    "whole", "fresh", "spot", "us", "ultra", "sensitive",
    # Anticoagulant/preservative additives named on the specimen-tube column
    # -- "Fluoride" for a glucose draw, "EDTA" for CBC/HbA1c -- printed
    # directly ahead of the test name exactly like "Serum"/"Plasma" already
    # were, just not previously listed. Without these, a test qualified
    # this way was rejected outright rather than matched: not a wrong
    # alias, no alias at all, since the prefix contained a word that was
    # not on this list. Confirmed on a real report where this was the
    # difference between three real tests being silently dropped entirely
    # and being recognised correctly.
    "fluoride", "edta", "heparin", "citrate", "oxalate",
    # "Approximate Mean Plasma Glucose" -- the same eAG test this project
    # already recognises under a dozen other phrasings, just with one more
    # qualifying word in front that was not on this list either. Confirmed
    # on a real report.
    "approximate", "estimated", "calculated",
}


def _find_analyte(text: str) -> tuple[Analyte, int] | None:
    """
    Match the longest analyte alias appearing near the start of a row.

    Restricted to the first part of the line because a range or a comment later
    in the row can easily contain a word that looks like another analyte. An
    exact match is tried first; only if that fails does it retry with look-alike
    characters folded, so a clean name never loses to a fuzzy one.
    """
    lowered = text.lower()
    head = lowered[:48]

    for alias, analyte in _ALIAS_INDEX:
        pos = head.find(alias)
        if pos == -1:
            continue
        before_ok = pos == 0 or not lowered[pos - 1].isalnum()
        after = pos + len(alias)
        after_ok = after >= len(lowered) or not lowered[after].isalnum()
        # A real word before the match usually means the alias is embedded
        # inside a longer, unrelated name rather than being the row's own
        # leading word -- "cholesterol" inside "NON HDL CHOLESTROL", "hdl"
        # inside "CORONARY RISK RATIO-1 (T.CHOL/HDL)". Two exceptions: a
        # specimen/timing qualifier ("Serum Total Cholesterol" is the same
        # test, not a different one), and a clean "(ALIAS)" parenthetical --
        # a lab's own abbreviation for the name just before it, e.g.
        # "Glycoslated Hb% Conc. (HbA1c)" -- distinguished from the ratio
        # case above by having nothing else inside the parentheses.
        prefix_raw = head[:pos]
        prefix_words = re.findall(r"[a-z]+", prefix_raw)
        # A prefix with no letters at all -- pure whitespace/indentation --
        # trivially satisfies "every word here is a qualifier" (there are
        # zero words to check), and that is fine: it is the same as no
        # prefix at all. But a *numbered list marker* ("1. HbA1C has been
        # endorsed by...") ALSO has no letters in its prefix, and treating
        # that the same way let a report's own footnote/methodology prose
        # match as if it were a real table row -- confirmed on a real
        # report, reproduced on a second, unrelated one, both fabricating a
        # bogus extra "reading" out of a sentence that just happened to
        # mention the test's name. A digit anywhere in the prefix is what
        # actually distinguishes the two: real qualifier words never
        # contain one.
        qualifier_ok = (all(w in _PREFIX_QUALIFIERS for w in prefix_words)
                        and not re.search(r"\d", prefix_raw))
        paren_ok = (pos > 0 and lowered[pos - 1] == "("
                    and after < len(lowered) and lowered[after] == ")")
        prefix_ok = pos == 0 or qualifier_ok or paren_ok
        if before_ok and after_ok and prefix_ok:
            # "hb" is the shortest, most generic alias this project has --
            # and OCR just as often misreads the "1" in "A1C" as a capital
            # "I", lowercase "i" or "l" as it reads it correctly. On those
            # reads, none of HbA1c's own aliases match exactly ("a1c" != a
            # literal "aic"/"alc"), the loop falls all the way down to bare
            # "hb", and returns immediately -- never reaching the loose/fold
            # fallback just below, which folds "1"/"l" the same way and
            # WOULD have matched correctly. Caught here instead of removing
            # "hb" outright, which stays a completely legitimate alias for
            # plain Haemoglobin the rest of the time. Confirmed on a real
            # report, reproduced identically across three of its four OCR
            # passes.
            if alias == "hb" and re.match(r"\s*a[i1l]\s*c\b", lowered[after:after + 6]):
                continue
            return analyte, after      # index is longest-first: first hit wins

    # Fallback: compare with look-alikes folded. Positions no longer line up
    # after folding, so the split point is recovered by walking the original
    # text until enough real characters have been consumed.
    folded_head = _loose(head)
    for alias, analyte in _LOOSE_INDEX:
        if len(alias) < 4 or not folded_head.startswith(alias):
            continue
        consumed = 0
        for i, ch in enumerate(text):
            if re.match(r"[A-Za-z0-9]", ch):
                consumed += 1
            if consumed == len(alias):
                return analyte, i + 1
        break
    return None


_BARE_RESULT_LINE = re.compile(r"^\(?\s*results?\s*:?\s*", re.IGNORECASE)
# A method note or two, plus an explanatory sentence, comfortably fits inside
# this; a genuinely unrelated section further down does not. Bounded on
# purpose so a heading with no continuation at all cannot wander into the
# next test's territory and misattribute its value.
_HEADING_LOOKAHEAD = 8


def _fold_heading_continuations(lines: list[Line]) -> list[Line]:
    """
    Bridge a name-only heading to a value written several lines below it.

    Some reports print the test name as its own row, then one or more lines
    of method or explanatory text, and only then the actual reading --
    "GLYCOSYLATED HAEMOGLOBIN" / "(Cation-Exchange Resin Method)" / "(...
    indicates blood sugar control...)" / "Result : 7.2 %". _parse_row only
    ever sees one line at a time, so the heading (a real alias match, but no
    number on it) and the bare "Result :" line (a number, but no analyte
    name of its own) both fail to parse on their own.

    Every line in between has to fail _find_analyte for the bridge to form,
    so this can only ever reach the value that actually belongs to the
    heading above it -- a real, different test starting in between stops
    the search rather than being skipped over. Confirmed on a real report.
    """
    out = list(lines)
    i = 0
    while i < len(out):
        text = out[i].text.strip()
        hit = _find_analyte(text) if len(text) >= 3 else None
        if hit and not _NUMBER.search(text[hit[1]:]):
            for j in range(i + 1, min(i + 1 + _HEADING_LOOKAHEAD, len(out))):
                cand_text = out[j].text.strip()
                if _find_analyte(cand_text) is not None:
                    break
                m = _BARE_RESULT_LINE.match(cand_text)
                if m and _NUMBER.search(cand_text[m.end():]):
                    out[i] = Line(segments=out[i].segments + out[j].segments)
                    del out[j]
                    break
        i += 1
    return out


def _looks_like_sex_continuation(line: Line) -> bool:
    """
    True for a line that carries the other half of a sex-split reference
    range printed on its own row -- "Female : 11.5 - 16.4)", or "(Method
    :Automated) Female : 12.0-15.5 Gms %)" when a method note shares the row
    -- the continuation of a "Haemoglobin ... (Male : 12.5-18.0" row one
    line up. Not anchored to the start: real reports interleave a trailing
    method annotation ahead of the sex word as often as not. Requires an
    actual number range, not just the bare word, and that the line not
    itself start a new test -- so a genuinely unrelated next row is never
    folded in by mistake. Confirmed on a real report.
    """
    text = line.text.strip()
    if not re.search(r"\b(?:men|male|women|female)\b", text, re.IGNORECASE):
        return False
    if not _RANGE_BETWEEN.search(text):
        return False
    return _find_analyte(text) is None


# "Ref. Range" (or "Reference Range") starting a line, on its own -- the
# reference range printed as its own row directly under the test's, rather
# than sharing the row with the result. Anchored to the start on purpose,
# unlike the sex-continuation label above: this exact phrase is common
# enough as an incidental mention elsewhere ("results should always be
# assessed in the context of the reference range") that matching it
# anywhere in a line would fold in far too much.
_RANGE_CONTINUATION_LABEL = re.compile(r"^\(?\s*ref(?:erence)?\.?\s*range\b",
                                       re.IGNORECASE)
# The two bounds on that line are as often written "60 100" (just a space)
# as "60-100" (a dash) -- a bare gap between two numbers is only safe to
# treat as a range at all because this is only ever tried on a line already
# confirmed to be nothing but a "Ref. Range" label plus its numbers.
_BARE_TWO_NUMBERS = re.compile(r"(\d+(?:[.,]\d+)?)\s+(\d+(?:[.,]\d+)?)")


def _looks_like_range_continuation(line: Line) -> bool:
    """
    True for a line that is nothing but a reference range printed on its own
    row directly below the test's name and result -- "Fasting Blood Sugar
    109 Mgs/dl" then, alone on the next line, "Ref. Range 60 100 Mgs/dl".
    Confirmed on a real report: three tests in a row each had this shape,
    and each one's own row genuinely carries no range at all -- so without
    folding the next line in, "no range printed" was not just misleading
    phrasing, it silently threw away a range that really was on the page,
    two lines down.
    """
    text = line.text.strip()
    if _find_analyte(text) is not None:
        return False
    if _RANGE_CONTINUATION_LABEL.match(text):
        return True
    # No explicit "Ref. Range" label at all -- some reports annotate the
    # range with just a population/method note instead ("Serum Euthyroid:
    # 0.80-2.00", "Adult: 4-6"). Accepted only when the line is short and
    # carries exactly one dash-separated range and no separate bound-style
    # number -- a real result row is never this sparse on its own, and a
    # genuinely unrelated range sitting on the very next line for some
    # other reason essentially never is either. Confirmed on a real
    # report: without this, a builtin fallback range replaced the report's
    # own -- 1.16 compared against a wrong 80-200 instead of the printed
    # 0.80-2.00 -- and got flagged as an implausible misread instead of
    # the correct, unremarkable normal result it actually was.
    return (len(text) <= 40 and len(_RANGE_BETWEEN.findall(text)) == 1
            and not _RANGE_BOUND.search(text))


# A tier's own label word says which direction it actually means, which the
# numbers alone cannot: "Diabetic: > 6.5" is a bound shaped exactly like an
# ordinary "normal means above this" reference, but landing in it is the
# ABNORMAL outcome, the opposite of what that bound shape usually implies.
# Checked in this order (high-direction first) only because "insufficient"
# contains "sufficient" as a literal substring, and must not be mistaken for
# the tier that word is stolen from.
# "High"/"Low" bare are added too -- the standard NCEP cholesterol tiering
# (Desirable / Borderline High / High) printed on a huge share of lipid
# panels uses exactly this wording, not "elevated"/"hyper-". Safe to include
# as bare words specifically because this whole check only ever runs inside
# _tier_from_line's already-narrow shape (a short line that is nothing but a
# label plus a bound/range) -- nowhere near broad enough for "high"/"low" in
# their ordinary generic sense to slip in from unrelated text.
_HIGH_ABNORMAL_WORDS = re.compile(
    r"\b(?:diabet\w*|elevated|excess\w*|hyper\w*|borderline|high)\b",
    re.IGNORECASE)
_LOW_ABNORMAL_WORDS = re.compile(
    r"\b(?:deficien\w*|insufficien\w*|hypo\w*|low)\b", re.IGNORECASE)


def _tier_from_line(text: str) -> tuple[float | None, float | None, str] | None:
    """A (low, high, abnormal) tier if this line is a bare bound/range
    annotation, else None. `abnormal` is "high"/"low" when the tier's own
    label says which direction landing in it means, else ""."""
    if len(text) > 45 or _find_analyte(text) is not None:
        return None
    abnormal = ("high" if _HIGH_ABNORMAL_WORDS.search(text)
                else "low" if _LOW_ABNORMAL_WORDS.search(text) else "")
    between = _RANGE_BETWEEN.search(text)
    if between:
        return _to_float(between.group(1)), _to_float(between.group(2)), abnormal
    bound = _RANGE_BOUND.search(text)
    if bound:
        limit = _to_float(bound.group(2))
        if bound.group(1).lower() in ("<", "<=", "upto", "up to", "less than"):
            return None, limit, abnormal
        return limit, None, abnormal
    return None


def _value_in_tier(numeric: float, tier: tuple[float | None, float | None, str]) -> bool:
    low, high, _abnormal = tier
    if low is None and high is None:
        return False
    if low is not None and numeric < low:
        return False
    if high is not None and numeric > high:
        return False
    return True


def _parse_range_continuation(line: Line) -> tuple[str, float | None, float | None] | None:
    """Pull (ref_text, low, high) out of a line already confirmed to be one."""
    text = _RANGE_CONTINUATION_LABEL.sub("", line.text.strip()).strip(" :.-")
    between = _RANGE_BETWEEN.search(text)
    if between:
        return between.group(0), _to_float(between.group(1)), _to_float(between.group(2))
    bound = _RANGE_BOUND.search(text)
    if bound:
        limit = _to_float(bound.group(2))
        if bound.group(1).lower() in ("<", "<=", "upto", "up to", "less than"):
            return bound.group(0), None, limit
        return bound.group(0), limit, None
    pair = _BARE_TWO_NUMBERS.search(text)
    if pair:
        return pair.group(0), _to_float(pair.group(1)), _to_float(pair.group(2))
    return None


def _parse_row(line: Line, sex: str | None = None,
               next_line: "Line | None" = None,
               next_line2: "Line | None" = None) -> AnalyteResult | None:
    """
    Turn one visual line of a lab table into a result, or None.

    `sex` comes from the report header. It is used only to choose between
    ranges that genuinely differ -- never to alter a value. `next_line` is
    the row immediately below, consulted for the one-sided-sex-range case
    below and (together with `next_line2`, the row after that) for a
    multi-tier reference range split across several lines -- everything
    else about this row still comes from `line` alone.
    """
    text = line.text.strip()
    if len(text) < 3:
        return None

    # Narrower than _NOT_A_ROW on purpose -- that full blocklist is only
    # ever applied in _parse_unknown_row, to a line with no matched analyte
    # at all, where rejecting broadly (even generic words like "desirable")
    # is safe. Here a real analyte has already matched, so the same full
    # list is too aggressive: "Desirable: <200" is completely ordinary
    # cholesterol reference-range phrasing, and blocking on it discarded a
    # genuine reading (confirmed -- this exact case regressed a previously-
    # passing real report the first time this used _NOT_A_ROW directly).
    # What actually needs catching is narrower: a diabetes-control-tier
    # label sharing a visual row with an analyte's own name ("(HbA1c)
    # Method:HPLC   6.1 - 7.0 % Good Control" -- a real layout, one
    # interpretation tier per line, aligned beside the label), which reads
    # as a second, wrong result for that analyte. Confirmed on a real report.
    if _CONTROL_TIER_ROW.search(text):
        return None

    # A square-bracketed *multi-word phrase* -- "[Primary Target of
    # Therapy]" -- is a reliable signal that this line is explanatory
    # legend text ("LDL Cholesterol (mg/dL) <100 Optimal, [Primary Target
    # of Therapy], 100-129 - Near Optimal...") rather than an actual result
    # row: several distinct tiers crammed onto one line, not the report's
    # own printed reading. Left alone, the first tier's own bound gets
    # mistaken for the row's value, with a later tier's range mistaken for
    # its reference. Requiring a space inside the brackets is what keeps
    # this from rejecting a real row that uses brackets the same way
    # parentheses are used elsewhere, for a short abbreviation -- "Estimated
    # Glucose Level [GH]" has no space inside its brackets and is
    # unaffected. Confirmed on a real report.
    if re.search(r"\[[^\]]*\s[^\]]*\]", text):
        return None

    hit = _find_analyte(text)
    if not hit:
        return None
    analyte, name_end = hit

    # An alias can match only the opening of a fuller parenthetical qualifier
    # -- "sugar (pp)" matching just the "(PP" of a printed "(PP - 1/2H)" --
    # because the loose look-alike fallback in _find_analyte walks alnum
    # characters rather than requiring the whole alias's punctuation to line
    # up. Left alone, the still-open "(" means the rest of that qualifier
    # ("- 1/2H)") lands in tail and its own digits get mistaken for the
    # result. Confirmed on a real report: "SUGAR (PP - 1/2H) : 142 mgs/dl"
    # read a result of 1, not 142, until this skipped past the close first.
    open_parens = text[:name_end].count("(") - text[:name_end].count(")")
    if open_parens > 0:
        close = text[name_end:].find(")")
        if close != -1:
            name_end += close + 1

    # A name-echoing parenthetical straight after the match -- "Cancer
    # Antigen 125 (CA-125)", "Estimated GFR (eGFR)" -- otherwise leaves its
    # own digits/letters sitting in tail, where a number inside it (the
    # "125" in "CA-125") gets mistaken for the row's actual result. Skipped
    # only when a "(...)" immediately follows the match, so a genuine value
    # or range starting right after the name is never touched. Confirmed on
    # a real report.
    echo = re.match(r"\s*\([A-Za-z0-9\-./]{2,15}\)", text[name_end:])
    if echo:
        name_end += echo.end()

    # A descriptive sentence about the test ("HbA1c is an index of your
    # blood sugar control for the past 3 months") starts with the test's
    # own name at position 0, which is unconditionally trusted -- unlike
    # the numbered-footnote case, there is no digit-prefix signal to catch
    # this one. A linking verb immediately after the name is a reliable,
    # narrow tell that this is prose describing the test, not a row
    # reporting it: a real result never reads "HbA1c is ...", it reads
    # "HbA1c 6.1 % ...". Confirmed on a real report, where the stray "3" in
    # "past 3 months" was otherwise mistaken for a fabricated result.
    if re.match(r"^\s+(?:is|was|are|were|measures?|indicates?|reflects?|"
                r"represents?|refers?)\b", text[name_end:], re.IGNORECASE):
        return None

    tail = text[name_end:]
    tail_original = tail

    # A report that splits its reference range by sex sometimes puts each
    # half on its own physical line rather than both on the row with the
    # test -- "Haemoglobin ... (Male : 12.5-18.0" then, alone on the very
    # next line, "Female : 11.5-16.4)". Only this row's own tail is ever
    # searched for the *result* itself (folded in below, not into `tail`),
    # so a continuation can only add a second reference range, never change
    # what value gets reported. Folding its text in here lets the sex-split
    # detection just below -- unchanged, already handles both halves on one
    # line -- see both halves the same way it always has. Only applies when
    # this row's own tail carries exactly one sex-labelled range: two would
    # mean the split is already complete, and zero means this was never a
    # split row to begin with. The pattern itself showed up on three real
    # reports; the fix was verified end to end (wrong sex-range in, right
    # one out) on one of them -- the other two are no longer available to
    # re-run, but the mechanism is the same layout, not a per-report guess.
    if (next_line is not None
            and re.search(r"\b(?:men|male|women|female)\b", tail_original,
                          re.IGNORECASE)
            and len(_RANGE_BETWEEN.findall(tail_original)) == 1
            and _looks_like_sex_continuation(next_line)):
        tail_original = tail_original + " " + next_line.text.strip()

    # Pull the reference range out first. It is the most distinctive thing on
    # the row, and removing it stops its numbers being mistaken for the result.
    ref_text, ref_low, ref_high = "", None, None
    sex_resolved = False
    between = _RANGE_BETWEEN.search(tail)
    bound = _RANGE_BOUND.search(tail)

    if between:
        ref_text = between.group(0)
        ref_low, ref_high = _to_float(between.group(1)), _to_float(between.group(2))
        tail = tail[:between.start()] + " " + tail[between.end():]
    elif bound:
        ref_text = bound.group(0)
        limit = _to_float(bound.group(2))
        if bound.group(1).lower() in ("<", "<=", "upto", "up to", "less than"):
            ref_high = limit
        else:
            ref_low = limit
        tail = tail[:bound.start()] + " " + tail[bound.end():]
    elif next_line is not None and _looks_like_range_continuation(next_line):
        # The result's own row carries no range at all -- it is printed as
        # its own row directly underneath instead. Only the range comes
        # from next_line; the result itself still comes entirely from this
        # row, same as the sex-split continuation above.
        found = _parse_range_continuation(next_line)
        if found:
            ref_text, ref_low, ref_high = found

    unit_match = _UNIT_RE.search(tail)
    unit = unit_match.group(1) if unit_match else ""
    if unit_match:
        tail = tail[:unit_match.start()] + " " + tail[unit_match.end():]

    # A result reported as "<0.005" or ">100" (below/above the assay's
    # detection limit) is bound-shaped exactly like a reference range is, and
    # when a *separate* between-style reference was already found above (the
    # common case: "<0.005 ... Adult : 0.27-5.35"), this leading bound is
    # still sitting untouched at the front of tail. Caught only when it is
    # the very first thing left there -- a bound anywhere else really is
    # leftover reference text, not the result -- and only when a distinct
    # reference range was already found, so a row with nothing but "<200"
    # (correctly claimed as the reference itself, above) never reaches this.
    # Confirmed on a real report: a TSH of "<0.005" was silently dropped from
    # the row entirely, because the blanket cleanup below erased it -- as if
    # it were leftover range noise -- before anything ever got a chance to
    # read it as the patient's actual result.
    leading_bound = re.match(r"^\s*(<=?|>=?)\s*(\d+(?:[.,]\d+)?)(?![A-Za-z])",
                              tail)
    if leading_bound and ref_text:
        value_text = leading_bound.group(0).strip()
        numeric = _to_float(leading_bound.group(2))
    else:
        # Remove *every* remaining range, not just the one taken as the
        # reference. A row like "SERUM HDL 44 MG% men 30-70 women 30-85"
        # carries two, and if the second survives, a bound from it gets
        # reported as the patient's result. Reporting a reference number as
        # a measurement is the worst thing this parser could do, so anything
        # range-shaped is cleared out first.
        tail = _RANGE_BETWEEN.sub(" ", tail)
        tail = _RANGE_BOUND.sub(" ", tail)

        number = _NUMBER.search(tail)
        if not number:
            return None
        value_text = number.group(0)
        numeric = _to_float(value_text)

    # A reference range split into several labelled tiers, each on its own
    # line ("Deficiency : <=20" / "Insufficiency: 21-29" / "Sufficiency:
    # >=30"), rather than one single bound. Only the first tier's own bound
    # was ever seen above (bound, not between -- a between-range is rarely
    # split into tiers this way, which keeps this narrowly scoped), so a
    # value that actually belongs in a later tier got compared against the
    # wrong one entirely: a healthy 34.9 checked against "<=20" (the
    # deficiency tier alone) came out flagged high. Gathered from up to two
    # more lines below and only switched to when the value falls inside
    # exactly one of them and it is not the tier already picked -- an
    # unambiguous case leaves the original, already-correct default alone.
    # Confirmed on a real report.
    tier_forced_flag = ""
    if numeric is not None and bound and not between:
        # The row's own tier needs the same abnormal-direction check as any
        # continuation tier -- "NON DIABETIC: < 5.7" is itself a tier in
        # exactly this same family, and correctly comes back "" (not
        # abnormal) since it is explicitly negated.
        own_abnormal = ("high" if _HIGH_ABNORMAL_WORDS.search(tail_original)
                        else "low" if _LOW_ABNORMAL_WORDS.search(tail_original)
                        else "")
        candidates: list[tuple[float | None, float | None, str]] = [
            (ref_low, ref_high, own_abnormal)]
        for nxt in (next_line, next_line2):
            if nxt is not None:
                tier = _tier_from_line(nxt.text.strip())
                if tier:
                    candidates.append(tier)
        if len(candidates) > 1:
            matching = [t for t in candidates if _value_in_tier(numeric, t)]
            if len(matching) == 1 and matching[0][:2] != (ref_low, ref_high):
                ref_low, ref_high, tier_forced_flag = matching[0]
                ref_text = (f"{ref_low:g} - {ref_high:g}"
                            if ref_low is not None and ref_high is not None
                            else f"< {ref_high:g}" if ref_high is not None
                            else f"> {ref_low:g}")

    # A row printing separate male and female ranges ("men 30-70 women 30-85")
    # cannot be judged from the row alone. Reporting the first range would give
    # a confident verdict against possibly the wrong sex, so the value is shown
    # with its ranges and no verdict at all.
    sex_split = bool(re.search(r"\b(?:men|male|women|female)\b", tail_original,
                               re.IGNORECASE)) and len(_RANGE_BETWEEN.findall(
                                   tail_original)) > 1

    # ...unless the header told us which sex applies. Then the matching half is
    # picked and the row is judged normally. Without a known sex it stays
    # unjudged: there the ambiguity is real, not a missing feature.
    if sex_split and sex:
        which = r"wom[ae]n|female" if sex == "F" else r"m[ae]n|male"
        picked = re.search(
            rf"(?:{which})\D{{0,12}}?(\d+(?:[.,]\d+)?)\s*(?:-|–|—|to)\s*"
            r"(\d+(?:[.,]\d+)?)",
            tail_original, re.IGNORECASE)
        if picked:
            ref_low = _to_float(picked.group(1))
            ref_high = _to_float(picked.group(2))
            ref_text = picked.group(0).strip()
            sex_split = False
            sex_resolved = True

    ref_source = "report"
    if ref_low is None and ref_high is None:
        ref_low, ref_high = analyte.range_for(sex)
        ref_source = "builtin"
        if ref_low is not None and ref_high is not None:
            ref_text = f"{ref_low:g} - {ref_high:g}"
        elif ref_high is not None:
            ref_text = f"< {ref_high:g}"
        elif ref_low is not None:
            ref_text = f"> {ref_low:g}"

    if sex_split:
        flag, implausible_note = "unknown", ""
    elif tier_forced_flag:
        # The winning tier's own bound is shaped like an ordinary "normal
        # means above/below this" reference, but its label says landing in
        # it IS the abnormal outcome ("Diabetic: > 6.5") -- the opposite of
        # what that bound shape would otherwise mean. _verdict is bypassed
        # entirely here rather than fed a bound it would misread.
        flag, implausible_note = tier_forced_flag, ""
    else:
        flag, implausible_note = _verdict(numeric, ref_low, ref_high)

    note = ""
    if implausible_note:
        note = implausible_note
    elif tier_forced_flag:
        note = ("the report labels this specific tier of its range as the "
                f"{'high' if tier_forced_flag == 'high' else 'low'} one, "
                "rather than printing a plain normal/abnormal boundary")
    elif sex_resolved:
        note = (f"the report prints ranges by sex; the "
                f"{'female' if sex == 'F' else 'male'} range was used")
    elif sex_split:
        ref_text = _clean_plain(tail_original[tail_original.lower().find(
            next(w for w in ("men", "male", "women", "female")
                 if w in tail_original.lower())):])[:60]
        note = ("the report prints separate ranges by sex -- not judged here, "
                "read the range that applies")
    elif ref_source == "builtin" and (ref_low is not None or ref_high is not None):
        note = ("no range printed on the report; compared against a typical adult"
                + (f" {'female' if sex == 'F' else 'male'} range" if sex and
                   (analyte.low_f is not None or analyte.high_f is not None)
                   else " range"))
        if analyte.note:
            note += f" ({analyte.note})"
    elif ref_low is None and ref_high is None:
        # Nothing printed and no built-in range: report the number, judge nothing.
        note = "no reference range printed and none assumed -- value shown as read"

    return AnalyteResult(
        key=analyte.key, name=analyte.name, panel=analyte.panel,
        value=value_text, numeric=numeric, unit=unit or analyte.unit,
        ref_text=ref_text.strip(), ref_low=ref_low, ref_high=ref_high,
        ref_source=ref_source, flag=flag,
        confidence=round(min(0.99, line.conf), 3),
        context=text[:120], note=note,
    )


# Lines that look like a test row but are really headers, footers or the
# interpretation tables labs print under HbA1c ("Non-Diabetes 4.0 to 6.0").
_NOT_A_ROW = re.compile(
    r"\b(?:page|report|end of|signature|verified|technician|pathologist|"
    r"consultant|address|phone|mobile|email|gstin|nabl|iso|bill no|"
    r"registration|printed|note|comment|method|interval|guidance|"
    r"good control|fair control|poor control|unsatisfactory|"
    # Stems, not whole words: the interpretation tables say "Non-Diabetes",
    # "Diabetic" and "Prediabetic", and a trailing \b would refuse to match
    # any of them because a letter follows the stem.
    r"non.?diabet\w*|diabet\w*|borderline|desirable|degree of control)",
    re.IGNORECASE)

# A narrower slice of the same idea, safe to apply even to a row whose name
# already matched a real analyte -- see _parse_row for why the full
# _NOT_A_ROW list above is not used there.
_CONTROL_TIER_ROW = re.compile(
    r"\b(?:good control|fair control|poor control|unsatisfactory|"
    r"degree of control)\b",
    re.IGNORECASE)


def parse_reference(ref_text: str) -> tuple[float | None, float | None, bool]:
    """
    Read a printed reference range back into bounds.

    Returns (low, high, ambiguous). `ambiguous` is True when the text carries
    more than one range and no sex was resolved -- those must not be judged.
    """
    if not ref_text:
        return None, None, False

    if (re.search(r"\b(?:men|male|women|female)\b", ref_text, re.IGNORECASE)
            and len(_RANGE_BETWEEN.findall(ref_text)) > 1):
        return None, None, True

    between = _RANGE_BETWEEN.search(ref_text)
    if between:
        return _to_float(between.group(1)), _to_float(between.group(2)), False

    bound = _RANGE_BOUND.search(ref_text)
    if bound:
        limit = _to_float(bound.group(2))
        if bound.group(1).lower() in ("<", "<=", "upto", "up to", "less than"):
            return None, limit, False
        return limit, None, False

    return None, None, False


# A result more than this many multiples beyond the reference bound is judged
# "check" rather than a plain high/low -- flagged as needing a human look
# rather than presented as a confident clinical reading. Confirmed on a real
# report: a printed 163.2 mg/dL cholesterol came back from OCR as 1632 (the
# decimal point was not recognised), an 8x inflation, and was shown as a plain
# HIGH result indistinguishable from a genuinely elevated one. Every real HIGH
# reading measured this session -- HbA1c 8.1 against a ref up to 5.6/6.4,
# glucose 133.4 and 180.9 against refs of 110/140, LDL 147 against 130 -- sat
# under 1.5x its bound, so 4x leaves a wide, safe margin between "genuinely
# elevated" and "very likely a misread digit or dropped decimal point."
_IMPLAUSIBLE_MULTIPLE = 4.0


def _verdict(numeric: float | None, low: float | None,
            high: float | None) -> tuple[str, str]:
    """
    The one place a value + range becomes a flag.

    Both the initial parse (_parse_row) and the export-time recompute
    (flag_for) call this, so a value edited after the fact is judged exactly
    the way it was judged the first time -- no separately-maintained copy of
    this logic to drift out of sync with this one.

    Returns (flag, note). flag is one of normal / high / low / check /
    unknown. "check" means implausible, not just outside range -- see
    _IMPLAUSIBLE_MULTIPLE above.
    """
    if numeric is None or (low is None and high is None):
        return "unknown", ""

    if high is not None and high > 0 and numeric > high * _IMPLAUSIBLE_MULTIPLE:
        return "check", (
            f"{numeric:g} is over {_IMPLAUSIBLE_MULTIPLE:g}x the upper end of "
            f"the range -- more likely a misread (a dropped decimal point, a "
            f"doubled digit) than a real result. Verify against the original "
            f"report before using this value.")
    if low is not None and low > 0 and numeric < low / _IMPLAUSIBLE_MULTIPLE:
        return "check", (
            f"{numeric:g} is under a quarter of the lower end of the range -- "
            f"more likely a misread than a real result. Verify against the "
            f"original report before using this value.")

    if high is not None and numeric > high:
        return "high", ""
    if low is not None and numeric < low:
        return "low", ""
    return "normal", ""


def flag_for(value_text: str, ref_text: str) -> str:
    """
    Recompute a verdict from a value and the range shown beside it.

    The export path calls this instead of believing the flag the browser sent.
    A user can edit a result after it was first judged, and a document that
    carries a stale HIGH next to a corrected number would be worse than one
    carrying no verdict at all.
    """
    numeric = _to_float(re.sub(r"[^\d.,\-]", "", value_text or ""))
    if numeric is None:
        return "unknown"

    low, high, ambiguous = parse_reference(ref_text)
    if ambiguous:
        return "unknown"
    flag, _note = _verdict(numeric, low, high)
    return flag


def _parse_unknown_row(line: Line) -> AnalyteResult | None:
    """
    Capture a table row whose test is not in ANALYTES.

    Every lab prints a slightly different menu, so a fixed dictionary will always
    miss something. Rather than drop those rows silently -- which would let a
    report look fully read when it was not -- anything row-shaped is surfaced
    under "Other rows", named exactly as printed and never flagged.
    """
    text = line.text.strip()
    if len(text) < 6 or _NOT_A_ROW.search(text):
        return None

    number = _NUMBER.search(text)
    if not number or number.start() == 0:
        return None                     # a row must be named before its value

    # The first number found can be part of the test's own printed name
    # rather than its result -- "Cancer Antigen 125 (CA-125)" repeats "125"
    # as the test's own abbreviation right after the name, and naively
    # taking the first number in the row reports that as a fabricated
    # result instead of the real value further along. Detected by the
    # number being echoed inside an immediately-following "(...)", and the
    # search retried past it for the row's actual value. Confirmed on a
    # real report, reproduced identically on a second, different lab.
    echo = re.match(r"\s*\([A-Za-z\-]*" + re.escape(number.group(0)) +
                    r"[A-Za-z\-]*\)", text[number.end():])
    if echo:
        retry = _NUMBER.search(text, number.end() + echo.end())
        if retry:
            number = retry

    name = text[:number.start()].strip(" :.-|\t")
    # Strip a trailing method column ("GOD-POD", "H.P.L.C") off the name.
    name = re.sub(r"\s+[A-Z][A-Z.\-]{2,}$", "", name).strip(" :.-")

    letters = sum(ch.isalpha() for ch in name)
    if letters < 3 or len(name) > 46:
        return None
    if not re.match(r"^[A-Za-z]", name):
        return None

    tail = text[number.end():]
    unit_match = _UNIT_RE.search(tail)
    ref = _RANGE_BETWEEN.search(tail) or _RANGE_BOUND.search(tail)

    # Require some corroboration that this really is a measurement row.
    if not unit_match and not ref:
        return None

    return AnalyteResult(
        key=f"other::{name.lower()}", name=name, panel="Other rows",
        value=number.group(0), numeric=_to_float(number.group(0)),
        unit=unit_match.group(1) if unit_match else "",
        ref_text=ref.group(0).strip() if ref else "",
        ref_source="report", flag="unknown",
        confidence=round(min(0.99, line.conf * 0.85), 3),
        context=text[:120],
        note="not a test this tool knows; shown exactly as printed, not judged",
    )


# Last-resort patient-name fallback for a report with no "Name" label
# anywhere at all -- see _find_name_before_age_sex_line.
_AGE_SEX_LINE = re.compile(
    r"^\(?\s*age\s*/?\s*sex\s*:|^\s*\d{1,3}\s*/\s*(?:M|F|Male|Female)\s*$",
    re.IGNORECASE)
_NAME_LIKE_LINE = re.compile(
    r"^(?:mr|mrs|ms|miss|mis|m/s|dr)\.?\s*[A-Za-z][A-Za-z.\s]{2,40}$",
    re.IGNORECASE)


def _find_name_before_age_sex_line(ocr: OcrResult) -> str:
    """
    A report that never labels the name field at all still reliably prints
    it as the line immediately above wherever "Age / Sex:" starts -- that's
    simply how the header block is laid out, label or not. Only ever tried
    once every labelled path in META_RULES has already come up empty, and
    only accepted when the candidate line itself looks like a name (an
    honorific, then letters, nothing else) -- so an unrelated line sitting
    above a coincidental "Age/Sex"-shaped bit of OCR noise elsewhere on the
    page is never mistaken for one. Confirmed on a real report.
    """
    for read in (ocr.variants or [ocr]):
        for i, line in enumerate(read.lines):
            if i == 0:
                continue
            if not _AGE_SEX_LINE.match(line.text.strip()):
                continue
            candidate = read.lines[i - 1].text.strip()
            if _NAME_LIKE_LINE.match(candidate):
                cleaned = _clean_name(candidate)
                if cleaned:
                    return cleaned
    return ""


def extract_report(ocr: OcrResult, filename: str = "") -> BloodReport:
    """Read patient metadata and every analyte row out of an OCR'd report."""
    meta_hits = core.extract_fields(ocr, enabled=(), custom_rules=META_RULES)
    meta = {key: hits[0].value for key, hits in meta_hits.items() if hits}
    if not meta.get("patient_name"):
        found_name = _find_name_before_age_sex_line(ocr)
        if found_name:
            meta["patient_name"] = found_name
    sex = detect_sex(meta)

    # Rows come from the single best read. Voting across variants does not help
    # here the way it does for an id -- a mangled row simply fails to parse and
    # the correct read of it wins by being the only one that parsed at all.
    reads = ocr.variants or [ocr]
    best_by_key: dict[tuple[str, int], AnalyteResult] = {}
    unknown: dict[str, AnalyteResult] = {}
    claimed_context: set[str] = set()

    for read in reads:
        # Rows are keyed by (analyte, how many times it has already appeared in
        # THIS read). Keying on the analyte alone would silently discard the
        # second fasting glucose on a report that lists two, while keying on the
        # value would split one row into several when variants misread it
        # differently. Occurrence order survives both.
        seen_in_read: Counter = Counter()
        folded = _fold_heading_continuations(read.lines)
        for i, line in enumerate(folded):
            next_line = folded[i + 1] if i + 1 < len(folded) else None
            next_line2 = folded[i + 2] if i + 2 < len(folded) else None
            row = _parse_row(line, sex, next_line, next_line2)
            if row is not None:
                occurrence = seen_in_read[row.key]
                seen_in_read[row.key] += 1
                if occurrence:
                    row.note = (row.note + "; " if row.note else "") + \
                        f"reading {occurrence + 1} of this test on the report"
                claimed_context.add(row.context)
                slot = (row.key, occurrence)
                current = best_by_key.get(slot)
                if current is None:
                    best_by_key[slot] = row
                else:
                    # Prefer a plausible reading over an implausible one
                    # regardless of which variant's OCR reported higher
                    # confidence -- different preprocessing passes
                    # sometimes read the very same printed digits
                    # differently (one pass prepends a stray extra digit,
                    # another does not), and a value flagged "check" (see
                    # _verdict) is exactly the signal that this particular
                    # read is probably the wrong one. Confirmed on a real
                    # report: one variant read a value correctly while
                    # three others each corrupted it differently, and
                    # without this the corrupted read's own higher OCR
                    # confidence was what won.
                    current_ok = current.flag != "check"
                    row_ok = row.flag != "check"
                    if row_ok and not current_ok:
                        best_by_key[slot] = row
                    elif row_ok == current_ok and row.confidence > current.confidence:
                        best_by_key[slot] = row

    # Second pass for rows no analyte claimed, so an unusual test still shows up.
    for read in reads:
        folded = _fold_heading_continuations(read.lines)
        for i, line in enumerate(folded):
            if line.text[:120] in claimed_context:
                continue
            next_line = folded[i + 1] if i + 1 < len(folded) else None
            next_line2 = folded[i + 2] if i + 2 < len(folded) else None
            if _parse_row(line, sex, next_line, next_line2) is not None:
                continue
            extra = _parse_unknown_row(line)
            if extra is None:
                continue
            current = unknown.get(extra.key)
            if current is None or extra.confidence > current.confidence:
                unknown[extra.key] = extra

    order = {a.key: i for i, a in enumerate(ANALYTES)}
    # Sort by analyte, then by which reading it was.
    results = [row for _, row in sorted(
        best_by_key.items(),
        key=lambda kv: (order.get(kv[0][0], 999), kv[0][1]))]

    # A variant that ran the words together (or dropped a space) produces a
    # different line than the one an analyte claimed, so the same measurement
    # can arrive twice: once recognised, once as "unknown." Compared against
    # the recognised row's own NAME here, this dedup silently failed on real
    # data -- an analyte's name is my chosen display label ("Glucose
    # (Fasting)"), which is never what the report itself printed ("Blood Sugar
    # (F)"), so the two could never match. What both reads actually share is
    # the raw line they came from, so that -- loose-folded, to absorb the
    # exact spacing difference between the two OCR passes -- is what gets
    # compared instead.
    taken_contexts = {_loose(r.context) for r in results if r.context}
    for extra in sorted(unknown.values(), key=lambda r: r.name.lower()):
        if extra.context and _loose(extra.context) in taken_contexts:
            continue
        results.append(extra)

    return BloodReport(
        filename=filename, meta=meta, results=results, full_text=ocr.text,
        full_text_display=ocr.display_text,
        engine=ocr.engine, ocr_confidence=ocr.mean_conf, elapsed=ocr.elapsed,
    )


# =========================================================================== #
#  5. EXPORT
# =========================================================================== #

def _timestamp() -> str:
    return datetime.now().strftime("%d %b %Y, %I:%M %p")


FLAG_WORD = {"high": "HIGH", "low": "LOW", "normal": "Normal",
            "check": "CHECK", "unknown": "-"}


def _rows(report: BloodReport) -> list[dict]:
    return [{
        "file": report.filename,
        "panel": r.panel,
        "test": r.name,
        "result": r.value,
        "unit": r.unit,
        "reference_range": r.ref_text,
        "range_source": r.ref_source,
        "flag": FLAG_WORD.get(r.flag, "-"),
        "note": r.note,
    } for r in report.results]


def build_csv(reports: Sequence[BloodReport]) -> bytes:
    columns = ["file", "panel", "test", "result", "unit", "reference_range",
               "range_source", "flag", "note"]
    out = io.StringIO(newline="")
    writer = csv.DictWriter(out, fieldnames=columns, extrasaction="ignore")
    writer.writeheader()
    for report in reports:
        writer.writerows(_rows(report))
    return out.getvalue().encode("utf-8-sig")


def build_json(reports: Sequence[BloodReport], include_raw_text: bool = True) -> bytes:
    payload = {
        "tool": REPORT_TITLE,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "disclaimer": DISCLAIMER,
        "report_count": len(reports),
        "reports": [],
    }
    for report in reports:
        entry = {
            "file": report.filename,
            "ocr_engine": report.engine,
            "ocr_confidence": round(report.ocr_confidence, 3),
            "patient": report.meta,
            "results": [{
                "test": r.name, "panel": r.panel, "result": r.value,
                "numeric": r.numeric, "unit": r.unit,
                "reference_range": r.ref_text, "range_source": r.ref_source,
                "flag": r.flag, "note": r.note,
            } for r in report.results],
            "out_of_range": [r.name for r in report.out_of_range],
        }
        if include_raw_text:
            entry["full_text"] = report.full_text
        payload["reports"].append(entry)
    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def build_txt(reports: Sequence[BloodReport], include_raw_text: bool = False) -> bytes:
    out = io.StringIO()
    out.write(f"{REPORT_TITLE}\n{'=' * len(REPORT_TITLE)}\n")
    out.write(f"Generated : {_timestamp()}\nReports   : {len(reports)}\n\n")
    out.write(DISCLAIMER + "\n")

    for idx, report in enumerate(reports, start=1):
        out.write(f"\n\n[{idx}] {report.filename}\n")
        out.write("-" * (len(report.filename) + 6) + "\n")
        for key in META_ORDER:
            if report.meta.get(key):
                label = next(r.label for r in META_RULES if r.key == key)
                out.write(f"  {label:<20}: {report.meta[key]}\n")

        out.write(f"\n  {'Test':<28}{'Result':>12} {'Unit':<12}"
                  f"{'Reference':<22}{'Flag':<8}\n")
        out.write("  " + "-" * 82 + "\n")
        for r in report.results:
            out.write(f"  {r.name:<28}{r.value:>12} {r.unit:<12}"
                      f"{r.ref_text:<22}{FLAG_WORD.get(r.flag, '-'):<8}\n")

        flagged = report.out_of_range
        out.write(f"\n  Outside the printed range: "
                  f"{', '.join(r.name for r in flagged) if flagged else 'none'}\n")

        if include_raw_text and report.full_text:
            out.write("\n  --- Full OCR text ---\n")
            for line in (report.full_text_display or report.full_text).splitlines():
                out.write(f"  {line}\n")

    return out.getvalue().encode("utf-8")


def build_docx(reports: Sequence[BloodReport], include_raw_text: bool = False) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    doc.add_heading(REPORT_TITLE, level=0)
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Generated: {_timestamp()}    "
                      f"Reports: {len(reports)}").italic = True

    note = doc.add_paragraph(DISCLAIMER)
    for run in note.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for idx, report in enumerate(reports, start=1):
        doc.add_heading(f"{idx}. {report.filename}", level=1)

        if report.meta:
            info = doc.add_table(rows=0, cols=2)
            info.style = "Light List Accent 1"
            for key in META_ORDER:
                if report.meta.get(key):
                    label = next(r.label for r in META_RULES if r.key == key)
                    cells = info.add_row().cells
                    cells[0].text = label
                    cells[1].text = report.meta[key]
            doc.add_paragraph()

        if not report.results:
            doc.add_paragraph("No test rows were recognised in this report.")
            continue

        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        for cell, heading in zip(table.rows[0].cells,
                                 ["Test", "Result", "Unit", "Reference", "Flag"]):
            cell.text = heading
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        for r in report.results:
            cells = table.add_row().cells
            cells[0].text = r.name
            cells[1].text = r.value
            cells[2].text = r.unit
            cells[3].text = r.ref_text
            cells[4].text = FLAG_WORD.get(r.flag, "-")
            for para in cells[1].paragraphs:
                for run in para.runs:
                    run.bold = True
            # "check" gets its own colour, distinct from clinical high/low --
            # it means "likely a misread," not "genuinely elevated."
            flag_colour = {"high": RGBColor(0xC0, 0x39, 0x2B),
                           "low": RGBColor(0xB7, 0x79, 0x0B),
                           "check": RGBColor(0x7B, 0x3F, 0xB2)}.get(r.flag)
            if flag_colour:
                for para in cells[4].paragraphs:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = flag_colour

        flagged = report.out_of_range
        summary = doc.add_paragraph()
        summary.add_run("Outside the printed range: ").bold = True
        summary.add_run(", ".join(r.name for r in flagged) if flagged else "none")

        borrowed = [r for r in report.results if r.ref_source == "builtin"]
        if borrowed:
            warn = doc.add_paragraph(
                f"{len(borrowed)} test(s) had no reference range printed on the "
                f"report; a typical adult range was used instead: "
                f"{', '.join(r.name for r in borrowed)}.")
            for run in warn.runs:
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x99, 0x66, 0x00)

        if include_raw_text and report.full_text:
            doc.add_heading("Full OCR text", level=2)
            block = doc.add_paragraph(report.full_text_display or report.full_text)
            for run in block.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_pdf(reports: Sequence[BloodReport], include_raw_text: bool = False) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            leftMargin=16 * mm, rightMargin=16 * mm,
                            topMargin=15 * mm, bottomMargin=15 * mm,
                            title=REPORT_TITLE)

    styles = getSampleStyleSheet()
    meta_style = ParagraphStyle("meta", parent=styles["Normal"], fontSize=8,
                                textColor=colors.HexColor("#666666"), leading=11)
    cell = ParagraphStyle("cell", parent=styles["Normal"], fontSize=8.5, leading=11)
    bold = ParagraphStyle("bold", parent=cell, fontName="Helvetica-Bold")
    mono = ParagraphStyle("mono", parent=styles["Normal"], fontName="Courier",
                          fontSize=7, leading=9)

    story = [
        Paragraph(REPORT_TITLE, styles["Title"]),
        Paragraph(f"Generated: {_timestamp()} &nbsp;|&nbsp; Reports: {len(reports)}",
                  meta_style),
        Spacer(1, 3 * mm),
        Paragraph(core._latin1_safe(DISCLAIMER), meta_style),
        Spacer(1, 6 * mm),
    ]

    for idx, report in enumerate(reports, start=1):
        story.append(Paragraph(f"{idx}. {escape(core._latin1_safe(report.filename))}",
                               styles["Heading2"]))

        if report.meta:
            bits = []
            for key in META_ORDER:
                if report.meta.get(key):
                    label = next(r.label for r in META_RULES if r.key == key)
                    bits.append(f"<b>{label}:</b> "
                                f"{escape(core._latin1_safe(report.meta[key]))}")
            story.append(Paragraph(" &nbsp;|&nbsp; ".join(bits), cell))
            story.append(Spacer(1, 3 * mm))

        if not report.results:
            story.append(Paragraph("No test rows were recognised.", cell))
            story.append(PageBreak() if idx < len(reports) else Spacer(1, 4 * mm))
            continue

        data = [[Paragraph(f"<b>{h}</b>", cell)
                 for h in ("Test", "Result", "Unit", "Reference", "Flag")]]
        flag_rows: list[tuple[int, str]] = []
        for n, r in enumerate(report.results, start=1):
            if r.flag in ("high", "low", "check"):
                flag_rows.append((n, r.flag))
            data.append([
                Paragraph(escape(core._latin1_safe(r.name)), cell),
                Paragraph(escape(core._latin1_safe(r.value)), bold),
                Paragraph(escape(core._latin1_safe(r.unit)), cell),
                Paragraph(escape(core._latin1_safe(r.ref_text)), cell),
                Paragraph(FLAG_WORD.get(r.flag, "-"), bold),
            ])

        table = Table(data, colWidths=[54 * mm, 24 * mm, 24 * mm, 42 * mm, 20 * mm],
                      repeatRows=1)
        style = [
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c8d0da")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f4f7fb")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]
        # "check" gets its own colour, distinct from clinical high/low -- it
        # means "likely a misread," not "genuinely elevated."
        flag_pdf_colour = {"high": colors.HexColor("#c0392b"),
                           "low": colors.HexColor("#b7790b"),
                           "check": colors.HexColor("#7b3fb2")}
        for row_no, flag in flag_rows:
            style.append(("TEXTCOLOR", (4, row_no), (4, row_no),
                          flag_pdf_colour[flag]))
        table.setStyle(TableStyle(style))
        story.append(table)

        flagged = report.out_of_range
        story.append(Spacer(1, 3 * mm))
        story.append(Paragraph(
            "<b>Outside the printed range:</b> " +
            (escape(core._latin1_safe(", ".join(r.name for r in flagged)))
             if flagged else "none"), cell))

        borrowed = [r for r in report.results if r.ref_source == "builtin"]
        if borrowed:
            story.append(Paragraph(
                f"{len(borrowed)} test(s) had no printed range; a typical adult "
                f"range was used: "
                f"{escape(core._latin1_safe(', '.join(r.name for r in borrowed)))}.",
                meta_style))

        if include_raw_text and report.full_text:
            story.append(Spacer(1, 3 * mm))
            story.append(Paragraph("Full OCR text", styles["Heading4"]))
            for line in core._latin1_safe(report.full_text_display or report.full_text).splitlines():
                if line.strip():
                    story.append(Paragraph(core._pdf_preserve_spaces(escape(line)), mono))

        if idx < len(reports):
            story.append(PageBreak())

    doc.build(story)
    return buffer.getvalue()


EXPORTERS = {
    "docx": (build_docx, "blood_report.docx",
             "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    "pdf": (build_pdf, "blood_report.pdf", "application/pdf"),
    "csv": (build_csv, "blood_report.csv", "text/csv"),
    "json": (build_json, "blood_report.json", "application/json"),
    "txt": (build_txt, "blood_report.txt", "text/plain"),
}


# =========================================================================== #
#  6. SELF TEST
# =========================================================================== #

SAMPLES = core.HERE / "samples"

SAMPLE_ROWS = [
    ("Haemoglobin", "11.2", "g/dL", "13.0 - 17.0"),
    ("Total Leucocyte Count", "9800", "/cumm", "4000 - 11000"),
    ("Platelet Count", "210000", "/cumm", "150000 - 450000"),
    ("Fasting Blood Sugar", "112", "mg/dL", "70 - 100"),
    ("HbA1c", "6.4", "%", "4.0 - 5.6"),
    ("Total Cholesterol", "186", "mg/dL", "< 200"),
    ("HDL Cholesterol", "38", "mg/dL", "> 40"),
    ("Triglycerides", "168", "mg/dL", "< 150"),
    ("Serum Creatinine", "0.9", "mg/dL", "0.6 - 1.3"),
    ("SGPT", "62", "U/L", "7 - 56"),
    ("TSH", "3.1", "uIU/mL", "0.4 - 4.0"),
    ("Vitamin D", "18.5", "ng/mL", "30 - 100"),
]

# Values chosen so the flags are unambiguous arithmetic.
EXPECTED_FLAGS = {
    "haemoglobin": "low", "wbc": "normal", "platelet": "normal",
    "glucose_f": "high", "hba1c": "high", "chol_total": "normal",
    "hdl": "low", "triglycerides": "high", "creatinine": "normal",
    "sgpt": "high", "tsh": "normal", "vit_d": "low",
}


def make_sample_report(path: Path) -> Path:
    """Draw a lab report resembling the layout Indian labs actually print."""
    from PIL import Image, ImageDraw

    width, height = 1000, 1080
    img = Image.new("RGB", (width, height), "#ffffff")
    d = ImageDraw.Draw(img)

    d.rectangle([0, 0, width, 96], fill="#12386b")
    d.text((40, 30), "SUNRISE DIAGNOSTICS", font=core._font(30, True), fill="#ffffff")

    info = [("Patient Name", "Meera Nair"), ("Age/Sex", "42 Y / F"),
            ("Patient ID", "SD2026081455"), ("Referred By", "Dr Anand Rao"),
            ("Collected On", "17 Aug 2026"), ("Reported On", "18 Aug 2026")]
    y = 124
    for label, value in info[:3]:
        d.text((40, y), f"{label}", font=core._font(19), fill="#6b7684")
        d.text((210, y), value, font=core._font(19, True), fill="#111820")
        y += 30
    y = 124
    for label, value in info[3:]:
        d.text((540, y), f"{label}", font=core._font(19), fill="#6b7684")
        d.text((700, y), value, font=core._font(19, True), fill="#111820")
        y += 30

    y = 244
    d.line([(40, y), (width - 40, y)], fill="#12386b", width=2)
    y += 14
    for head, x in [("TEST", 40), ("RESULT", 470), ("UNIT", 620), ("REFERENCE", 760)]:
        d.text((x, y), head, font=core._font(18, True), fill="#12386b")
    y += 30
    d.line([(40, y), (width - 40, y)], fill="#c8d0da", width=1)
    y += 14

    for name, value, unit, ref in SAMPLE_ROWS:
        d.text((40, y), name, font=core._font(20), fill="#111820")
        d.text((470, y), value, font=core._font(20, True), fill="#111820")
        d.text((620, y), unit, font=core._font(20), fill="#4a5460")
        d.text((760, y), ref, font=core._font(20), fill="#4a5460")
        y += 40

    d.text((40, y + 24), "** End of Report **", font=core._font(18), fill="#8a94a6")

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return path


def run_selftest() -> int:
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    if not core.available_engines():
        print("FAIL: no OCR engine. pip install rapidocr-onnxruntime")
        return 1

    sample = make_sample_report(SAMPLES / "sample_bloodtest.png")
    print(f"Sample report written : {sample}")

    ocr = core.read_image(sample, accuracy="fast")
    report = extract_report(ocr, filename=sample.name)
    print(f"OCR engine            : {report.engine}  "
          f"confidence {report.ocr_confidence:.0%}")

    print("\n--- Patient details ---")
    for key in META_ORDER:
        if report.meta.get(key):
            label = next(r.label for r in META_RULES if r.key == key)
            print(f"  {label:<20}: {report.meta[key]}")

    print(f"\n--- Test rows ({len(report.results)} of {len(SAMPLE_ROWS)}) ---")
    print(f"  {'Test':<30}{'Result':>10} {'Unit':<10}{'Reference':<18}"
          f"{'Flag':<8}{'Range from'}")
    for r in report.results:
        print(f"  {r.name:<30}{r.value:>10} {r.unit:<10}{r.ref_text:<18}"
              f"{FLAG_WORD.get(r.flag, '-'):<8}{r.ref_source}")

    print("\n--- Export formats ---")
    for ext, (builder, _, _) in EXPORTERS.items():
        blob = builder([report]) if ext == "csv" else builder([report], False)
        out = SAMPLES / f"sample_bloodtest.{ext}"
        out.write_bytes(blob)
        print(f"  {ext:<5} {len(blob):>8,} bytes -> {out.name}")

    print("\n--- Flag check (arithmetic against the printed range) ---")
    found = {r.key: r.flag for r in report.results}
    wrong, missing = [], []
    for key, expected in EXPECTED_FLAGS.items():
        name = ANALYTES_BY_KEY[key].name
        if key not in found:
            missing.append(name)
            print(f"  {name:<30} NOT FOUND (expected {expected})")
        elif found[key] != expected:
            wrong.append(name)
            print(f"  {name:<30} got {found[key]:<8} expected {expected}  MISMATCH")
        else:
            print(f"  {name:<30} {found[key]:<8} OK")

    recall = (len(EXPECTED_FLAGS) - len(missing)) / len(EXPECTED_FLAGS)
    print(f"\nrows found: {len(EXPECTED_FLAGS) - len(missing)}/{len(EXPECTED_FLAGS)} "
          f"({recall:.0%})   wrong flags: {len(wrong)}")

    if wrong or recall < 0.8:
        print("\nFAIL")
        return 1
    print("\nPASS: rows parsed, ranges read from the report, flags correct.")
    return 0


if __name__ == "__main__":
    sys.exit(run_selftest())
