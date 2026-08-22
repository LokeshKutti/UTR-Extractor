"""
UTR / Payment Detail Extractor -- core engine
=============================================
OCR, field extraction and document export. No web framework in here: the HTTP
layer lives in server.py and the interface in web/index.html.

    python core.py        run the end-to-end self test

Four sections:

    1. OCR            image cleanup, pluggable engines, line grouping
    2. EXTRACTION     field rules and the four matching strategies
    3. EXPORT         DOCX / PDF / CSV / JSON / TXT builders
    4. SELF TEST      end-to-end smoke test on a generated receipt
"""

from __future__ import annotations

import csv
import io
import json
import re
import sys
import threading
from collections import Counter, defaultdict
from dataclasses import dataclass, field, replace
from datetime import datetime
from html import escape
from pathlib import Path
from statistics import median
from typing import Callable, Iterable, Sequence

import numpy as np
from PIL import Image, ImageEnhance, ImageOps

APP_TITLE = "UTR / Payment Detail Extractor"

# Under PyInstaller, __file__ for a bundled module resolves inside the frozen
# archive, not a real directory -- sibling data files like web/ do not live
# there. sys._MEIPASS is PyInstaller's own answer to "where did you actually
# put the data files": for a onefile build it is the temp extraction dir, and
# for a onedir build (used here) it is the _internal folder next to the exe --
# a location that first looked like a plain sys.executable.parent guess would
# get wrong, since PyInstaller 6.x moved bundled data into that subfolder
# rather than placing it beside the exe itself.
HERE = Path(getattr(sys, "_MEIPASS", Path(sys.executable).parent)) \
    if getattr(sys, "frozen", False) else Path(__file__).parent


# =========================================================================== #
#  1. OCR
# =========================================================================== #

# Engines are imported lazily so a missing optional backend never breaks import.
_ENGINE_CACHE: dict[str, object] = {}

# RapidOCR writes its per-call thresholds onto the shared engine object
# (self.text_detector.postprocess_op.box_thresh = ...) before running. Two
# threads inside that call would overwrite each other's settings, and since each
# variant deliberately passes different thresholds the damage would be silent
# wrong output rather than a crash. OCR is CPU-bound anyway, so serialising it
# costs nothing that parallelism would have won -- the point of running it off
# the event loop is that the *server* stays responsive, not that OCR overlaps.
_OCR_LOCK = threading.Lock()


@dataclass
class Segment:
    """One OCR text box."""
    text: str
    conf: float
    x: float          # left edge
    y: float          # vertical centre
    height: float


@dataclass
class Line:
    """Segments that share a visual line, ordered left to right."""
    segments: list[Segment] = field(default_factory=list)

    @property
    def text(self) -> str:
        return " ".join(s.text for s in self.segments).strip()

    @property
    def conf(self) -> float:
        if not self.segments:
            return 0.0
        return sum(s.conf for s in self.segments) / len(self.segments)


@dataclass
class OcrResult:
    lines: list[Line] = field(default_factory=list)
    engine: str = ""
    elapsed: float = 0.0
    # Name of the preprocessing variant that produced this read.
    variant: str = ""
    # Above accuracy "fast", every variant's read is kept here so the extractor
    # can compare them and vote. The outer OcrResult stays the single best read,
    # which is what gets shown as "full OCR text".
    variants: list["OcrResult"] = field(default_factory=list)
    # Set only when accuracy="auto". `quality` is the blur/resolution read on the
    # image before OCR ran; `accuracy_used` is the tier that was actually run
    # (which may be higher than the image's own quality implied, if the first
    # pass came back too thin and the safety net escalated); `escalated` marks
    # that second case specifically, so the interface can say so.
    quality: dict = field(default_factory=dict)
    accuracy_used: str = ""
    escalated: bool = False

    @property
    def text(self) -> str:
        return "\n".join(ln.text for ln in self.lines if ln.text)

    @property
    def mean_conf(self) -> float:
        confs = [s.conf for ln in self.lines for s in ln.segments]
        return sum(confs) / len(confs) if confs else 0.0

    @property
    def word_count(self) -> int:
        return len(self.text.split())


class NoEngineAvailable(RuntimeError):
    pass


def load_image(data: bytes | str | Path | Image.Image) -> Image.Image:
    """Accept raw bytes, a path, or a PIL image; always return RGB."""
    if isinstance(data, Image.Image):
        img = data
    elif isinstance(data, (bytes, bytearray)):
        img = Image.open(io.BytesIO(data))
    else:
        img = Image.open(data)
    img = ImageOps.exif_transpose(img)      # honour phone-camera rotation
    return img.convert("RGB")


MAX_PDF_PAGES = 15
PDF_RENDER_DPI = 200        # enough for 8pt lab-report type; 300 is 2x the cost


def is_pdf(data: bytes | str | Path) -> bool:
    if isinstance(data, (bytes, bytearray)):
        return data[:5] == b"%PDF-"
    try:
        return Path(data).suffix.lower() == ".pdf"
    except (TypeError, ValueError):
        return False


# A page with fewer real words than this is treated as having no usable text
# layer -- a scanned or photographed page wrapped in a PDF typically extracts
# to exactly zero words, so this only needs to clear "obviously not empty,"
# not distinguish a short report from a long one.
_PDF_NATIVE_MIN_WORDS = 10


def _pdf_native_lines(page) -> list["Line"] | None:
    """
    Read one PDF page's embedded text directly, skipping OCR entirely.

    A PDF generated by a real system -- a hospital's report software, a
    payment gateway's receipt template -- carries the exact text it was built
    from, not a picture of it. Rendering that to an image and then OCR'ing the
    image throws away a source of ground truth and replaces it with a guess:
    confirmed on a real 8-page lab report, every transcription error found in
    it (a TSH of 1.58 misread as 3, a test name colliding with an unrelated
    one, a sex-split range failing to resolve) traced back to OCR/layout
    imprecision that this path cannot suffer from, because there is no
    recognition step -- the text is copied, not read.

    Returns None when the page looks like a scanned image with no real text
    layer (a camera photo or scan wrapped in a PDF container), so the caller
    falls back to rendering + OCR for that one page. A mixed document -- some
    native pages, some scanned -- gets each page handled the way it actually
    needs.
    """
    words = page.get_text("words")     # (x0, y0, x1, y1, text, block, line, word_no)
    segments = [
        Segment(text=w[4], conf=1.0, x=w[0], y=(w[1] + w[3]) / 2, height=w[3] - w[1])
        for w in words if w[4].strip()
    ]
    if len(segments) < _PDF_NATIVE_MIN_WORDS:
        return None
    return _group_lines(segments)


def _open_pdf(data: bytes | str | Path):
    """The PyMuPDF document object for either a path or raw bytes."""
    try:
        # PyMuPDF renamed its module; "fitz" still works but warns on every use.
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError as exc:
            raise RuntimeError(
                "PDF support needs PyMuPDF. Install it with:  pip install pymupdf"
            ) from exc

    if isinstance(data, (bytes, bytearray)):
        return fitz.open(stream=bytes(data), filetype="pdf")
    return fitz.open(str(data))


def load_pages(data: bytes | str | Path | Image.Image) -> list[Image.Image]:
    """
    Every page of a document as an image.

    Images give a single page. PDFs are rendered at PDF_RENDER_DPI -- lab
    reports and receipts arrive as PDFs at least as often as photos, and a
    text-layer PDF renders far cleaner than any camera shot of the same page.

    This always rasterises, even for a page read_document() will end up
    reading natively instead -- it exists for callers that only want a
    picture of the document (the thumbnail preview), not its text.
    """
    if isinstance(data, Image.Image) or not is_pdf(data):
        return [load_image(data)]

    pages: list[Image.Image] = []
    with _open_pdf(data) as doc:
        for index, page in enumerate(doc):
            if index >= MAX_PDF_PAGES:
                break
            pix = page.get_pixmap(dpi=PDF_RENDER_DPI)
            pages.append(Image.frombytes("RGB", (pix.width, pix.height), pix.samples))
    if not pages:
        raise RuntimeError("That PDF has no pages that could be rendered.")
    return pages


def preprocess(img: Image.Image, mode: str = "auto") -> Image.Image:
    """
    Clean an image up for OCR.

      none       -- pass through untouched
      auto       -- upscale small screenshots, mild contrast + sharpen
      aggressive -- auto, plus grayscale and a binarising threshold

    Payment screenshots are usually small and low-contrast, so upscaling before
    OCR is what most often turns a mangled UTR into a clean one.
    """
    if mode == "none":
        return img

    out = img

    # Upscale until the short edge is at least 1000px (capped at 3x).
    short_edge = min(out.size)
    if short_edge < 1000:
        scale = min(3.0, 1000 / max(short_edge, 1))
        out = out.resize((int(out.width * scale), int(out.height * scale)), Image.LANCZOS)

    out = ImageEnhance.Contrast(out).enhance(1.4)
    out = ImageEnhance.Sharpness(out).enhance(1.6)

    if mode == "aggressive":
        # Otsu picks the cut from this image's own histogram, so a dim photo and
        # a bright scan each get the threshold they need.
        out = binarise(ImageOps.autocontrast(ImageOps.grayscale(out), cutoff=2)
                       .convert("RGB"))

    return out


# --------------------------------------------------------------------------- #
# Camera-photo correction
# --------------------------------------------------------------------------- #
#
# A screenshot arrives flat, square and evenly lit. A phone photo of a printed
# report arrives rotated a few degrees, lit from one side, and often creased.
# These three steps are no-ops on a screenshot and the difference between
# readable and not on a photo.

def _skew_angle(img: Image.Image, limit: float = 8.0) -> float:
    """
    Estimate page rotation by projection profile.

    Text rows produce sharp peaks and troughs in a row-sum profile only when
    they are horizontal, so the angle whose profile varies most is the angle
    that straightens the page. Estimated on a downscaled copy -- the angle is a
    property of the layout, not of the resolution, and full size would cost far
    more for no extra precision.
    """
    small = ImageOps.grayscale(img)
    if small.width > 800:
        small = small.resize((800, max(1, int(small.height * 800 / small.width))))
    data = 255 - np.asarray(small, dtype=np.float32)      # ink = high values
    if data.std() < 5:                                    # effectively blank
        return 0.0

    import cv2
    height, width = data.shape
    centre = (width / 2, height / 2)
    best_angle, best_score = 0.0, -1.0
    for angle in np.arange(-limit, limit + 0.5, 0.5):
        matrix = cv2.getRotationMatrix2D(centre, float(angle), 1.0)
        rotated = cv2.warpAffine(data, matrix, (width, height),
                                 flags=cv2.INTER_LINEAR, borderValue=0)
        profile = rotated.sum(axis=1)
        score = float(((profile[1:] - profile[:-1]) ** 2).sum())
        if score > best_score:
            best_score, best_angle = score, float(angle)
    return best_angle


def deskew(img: Image.Image, threshold: float = 0.6) -> Image.Image:
    """Straighten a photographed page. Left alone if it is already square."""
    angle = _skew_angle(img)
    if abs(angle) < threshold:
        return img
    return img.rotate(angle, resample=Image.BICUBIC, expand=True,
                      fillcolor=(255, 255, 255))


def flatten_illumination(img: Image.Image) -> Image.Image:
    """
    Remove the lighting gradient a phone camera leaves behind.

    Dividing the page by a heavily blurred copy of itself cancels whatever is
    varying slowly -- shadow, lamp falloff, the curve of a folded sheet -- while
    leaving the sharp ink alone. Without this, one global threshold cannot serve
    both the bright and the shadowed half of the same photo.
    """
    import cv2
    gray = np.asarray(ImageOps.grayscale(img), dtype=np.float32)
    sigma = max(img.width, img.height) / 30.0
    background = cv2.GaussianBlur(gray, (0, 0), sigmaX=sigma, sigmaY=sigma)
    flat = np.clip(gray / np.maximum(background, 1.0) * 200.0, 0, 255)
    return Image.fromarray(flat.astype(np.uint8)).convert("RGB")


def binarise(img: Image.Image) -> Image.Image:
    """
    Otsu threshold -- chosen from the image, not hardcoded.

    The old fixed cut at 140 assumed a bright scan; on a dim photo it turned the
    whole page black, and on a bright one it erased faint print.
    """
    import cv2
    gray = np.asarray(ImageOps.grayscale(img))
    _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY | cv2.THRESH_OTSU)
    return Image.fromarray(binary).convert("RGB")


# --------------------------------------------------------------------------- #
# Image quality -- deciding how hard to work before OCR even runs
# --------------------------------------------------------------------------- #
#
# A phone photo of a report is not always as clean as a screenshot. Rather than
# make the user notice that and switch to "Maximum" by hand, accuracy="auto"
# looks at the image first and picks the effort level itself.

TIER_ORDER = ["fast", "balanced", "high"]

# Laplacian variance (a standard focus measure -- the second derivative of the
# image, which is large at crisp edges and collapses toward zero as an edge
# blurs out) measured after scaling every image to the same reference size.
# Unscaled, the same measure scores a small sharp image *higher* than a large
# sharp one purely because of pixel density, which would misjudge exactly the
# small, far-away camera photos this is meant to catch -- so every image is
# resampled to this reference short edge before the score is taken.
_QUALITY_REF_EDGE = 900

# Chosen by measuring this metric across a clear render, several blur levels,
# reduced-contrast copies and a genuinely under-resolved image, then reading
# off where the values group (see the project's dev notes): everything actually
# clear scored in the hundreds or higher, real blur fell under ~70, and heavy
# blur or a too-small original fell under ~15. There is a wide gap on both
# sides, so exact placement within it is not sensitive.
_QUALITY_POOR_BELOW = 15.0
_QUALITY_SOFT_BELOW = 70.0

# If the read comes back this thin, something went wrong worth trying harder
# for. Measured clean images sit around 0.80 confidence, so 0.55 has a wide
# margin before it can fire on a good scan.
#
# A word-count floor was tried too and dropped: a legitimately small, cropped
# image (just "UTR No. 884512336790", say) will always have few words to find
# regardless of how clear it is -- that is content sparsity, not a quality
# problem, and measured at 0.858 confidence on three words, well above the
# threshold that actually matters. An empty read (zero words) is the one
# word-count reading that stays unambiguous regardless of how sparse the
# source is: nothing was read at all, which is always worth another attempt.
_ESCALATE_BELOW_CONF = 0.55


def estimate_quality(img: Image.Image, ref: int = _QUALITY_REF_EDGE) -> dict:
    """
    Score an image for blur / effective resolution before OCR runs.

    Returns {laplacian, short_edge, tier, reasons}. `tier` is "clear", "soft"
    or "poor". Contrast was tried as a second signal and dropped: a plain
    report is mostly white page, so whole-image contrast measures came out low
    even on a perfectly crisp scan -- they were measuring how much of the page
    is background, not whether the text on it is legible. Laplacian variance on
    a normalised copy does not have that problem, since it is already local to
    edges rather than averaged across the empty page around them.
    """
    import cv2

    short_edge = min(img.size)
    scale = ref / max(short_edge, 1)
    normed = img.resize((max(1, int(img.width * scale)), max(1, int(img.height * scale))),
                        Image.LANCZOS)
    gray = np.asarray(ImageOps.grayscale(normed))
    laplacian = float(cv2.Laplacian(gray, cv2.CV_64F).var())

    if laplacian < _QUALITY_POOR_BELOW:
        tier, reasons = "poor", ["looks blurry or was too small to read clearly"]
    elif laplacian < _QUALITY_SOFT_BELOW:
        tier, reasons = "soft", ["looks a little soft"]
    else:
        tier, reasons = "clear", []

    return {"laplacian": round(laplacian, 1), "short_edge": short_edge,
            "tier": tier, "reasons": reasons}


def _upscale(img: Image.Image, target_short: int, cap: float) -> Image.Image:
    """Grow an image until its short edge reaches target_short."""
    short = min(img.size)
    if short >= target_short:
        return img
    scale = min(cap, target_short / max(short, 1))
    return img.resize((int(img.width * scale), int(img.height * scale)), Image.LANCZOS)


def _enhance(img: Image.Image, contrast: float = 1.4, sharp: float = 1.6) -> Image.Image:
    out = ImageEnhance.Contrast(img).enhance(contrast)
    return ImageEnhance.Sharpness(out).enhance(sharp)


def build_variants(
    img: Image.Image, preprocess_mode: str, accuracy: str
) -> list[tuple[str, Image.Image, dict]]:
    """
    Produce the differently-prepared copies of an image to read.

    Each entry is (name, image, per-call OCR options). The variants are chosen
    to fail *differently* rather than to all be slightly better -- a blur that
    defeats the sharpened pass often leaves the binarised pass intact, and it is
    that disagreement the voting step exploits.
    """
    # Straightening happens once, before anything is derived from the image, so
    # every variant reads the same square page. On a screenshot the estimated
    # angle is ~0 and this returns the original untouched.
    img = deskew(img)

    primary = preprocess(img, preprocess_mode)
    variants: list[tuple[str, Image.Image, dict]] = [("clean", primary, {})]

    if accuracy == "fast":
        if preprocess_mode == "auto":
            variants.append(("raw", img, {}))
        return variants

    variants.append(("raw", img, {}))

    # A much larger render, read with looser detection so faint thin strokes
    # still form boxes. This is the single most useful extra pass on the small
    # screenshots people actually paste in.
    variants.append((
        "big",
        _enhance(_upscale(img, 1800, 4.0), 1.5, 1.9),
        {"text_score": 0.4, "unclip_ratio": 1.9, "box_thresh": 0.4},
    ))
    # Same picture, permissive thresholds: recovers characters the default
    # thresholds drop, at the cost of more noise (which voting then discards).
    variants.append((
        "loose", primary,
        {"text_score": 0.3, "unclip_ratio": 2.1, "box_thresh": 0.32},
    ))

    if accuracy != "high":
        return variants

    # Kept at 1500px on purpose. OCR cost scales with pixel area, so a 2600px
    # render costs ~3x a 1500px one; measured on this machine the extra scale
    # bought nothing an existing pass had not already read, while pushing a
    # single image past 40s. Two cheap passes that fail differently beat one
    # expensive pass that fails the same way as "big".
    base = _upscale(img, 1500, 4.0)
    gray = ImageOps.autocontrast(ImageOps.grayscale(base), cutoff=1)
    variants.append(("gray", gray.convert("RGB"), {"text_score": 0.35}))
    variants.append((
        "binary", binarise(gray.convert("RGB")),
        {"text_score": 0.35, "unclip_ratio": 1.8},
    ))
    # For photographs: cancel the lighting gradient, then threshold. Useless on
    # a screenshot, but it is the pass that rescues a shadowed paper report.
    variants.append((
        "camera", binarise(flatten_illumination(base)),
        {"text_score": 0.35, "unclip_ratio": 1.9, "box_thresh": 0.38},
    ))
    return variants


def _group_lines(segments: Sequence[Segment]) -> list[Line]:
    """Merge segments whose vertical centres are close into single lines."""
    if not segments:
        return []

    heights = [s.height for s in segments if s.height > 0]
    tolerance = (median(heights) * 0.6) if heights else 10.0

    ordered = sorted(segments, key=lambda s: (s.y, s.x))
    lines: list[Line] = []
    current: list[Segment] = [ordered[0]]

    for seg in ordered[1:]:
        anchor = sum(s.y for s in current) / len(current)
        if abs(seg.y - anchor) <= tolerance:
            current.append(seg)
        else:
            lines.append(Line(sorted(current, key=lambda s: s.x)))
            current = [seg]

    lines.append(Line(sorted(current, key=lambda s: s.x)))
    return lines


def _box_metrics(box) -> tuple[float, float, float]:
    """Return (left, vertical-centre, height) for a 4-point polygon."""
    pts = np.asarray(box, dtype=float).reshape(-1, 2)
    xs, ys = pts[:, 0], pts[:, 1]
    return float(xs.min()), float((ys.min() + ys.max()) / 2), float(ys.max() - ys.min())


def available_engines() -> list[str]:
    """Which OCR backends can actually run on this machine, best first."""
    found = []
    try:
        import rapidocr_onnxruntime  # noqa: F401
        found.append("rapidocr")
    except Exception:
        pass
    try:
        import easyocr  # noqa: F401
        found.append("easyocr")
    except Exception:
        pass
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        found.append("tesseract")
    except Exception:
        pass
    return found


def _run_rapidocr(img: Image.Image, **opts) -> tuple[list[Segment], float]:
    from rapidocr_onnxruntime import RapidOCR

    # The lock covers construction *and* the call: the engine keeps per-call
    # thresholds on itself, so the settings and the run they belong to must not
    # be separated by another thread.
    with _OCR_LOCK:
        engine = _ENGINE_CACHE.get("rapidocr")
        if engine is None:
            # Constructed with no overrides on purpose. Detector kwargs are broken
            # in rapidocr-onnxruntime 1.2.3 (update_det_params reads model_path
            # unconditionally and raises KeyError), and they are not needed anyway:
            # limit_type "min" only ever scales an image *up* to 736, so the larger
            # LANCZOS renders that build_variants produces pass through untouched.
            engine = RapidOCR()
            _ENGINE_CACHE["rapidocr"] = engine

        # box_thresh / unclip_ratio / text_score are honoured per call, letting
        # each variant trade recall against precision.
        result, elapse = engine(np.asarray(img), **opts)
    segments: list[Segment] = []
    for box, text, score in (result or []):
        if not str(text).strip():
            continue
        left, ycentre, height = _box_metrics(box)
        segments.append(Segment(str(text).strip(), float(score), left, ycentre, height))

    took = float(elapse[0]) if isinstance(elapse, (list, tuple)) and elapse else 0.0
    return segments, took


def _run_easyocr(img: Image.Image, **opts) -> tuple[list[Segment], float]:
    import easyocr

    with _OCR_LOCK:
        reader = _ENGINE_CACHE.get("easyocr")
        if reader is None:
            reader = easyocr.Reader(["en"], gpu=False, verbose=False)
            _ENGINE_CACHE["easyocr"] = reader
        detections = reader.readtext(np.asarray(img))

    segments: list[Segment] = []
    for box, text, score in detections:
        if not str(text).strip():
            continue
        left, ycentre, height = _box_metrics(box)
        segments.append(Segment(str(text).strip(), float(score), left, ycentre, height))
    return segments, 0.0


def _run_tesseract(img: Image.Image, **opts) -> tuple[list[Segment], float]:
    import pytesseract
    from pytesseract import Output

    with _OCR_LOCK:
        data = pytesseract.image_to_data(img, output_type=Output.DICT)
    segments: list[Segment] = []
    for i, word in enumerate(data["text"]):
        word = word.strip()
        conf = float(data["conf"][i])
        if not word or conf < 0:
            continue
        top, height = float(data["top"][i]), float(data["height"][i])
        segments.append(
            Segment(word, conf / 100.0, float(data["left"][i]), top + height / 2, height)
        )
    return segments, 0.0


_RUNNERS = {
    "rapidocr": _run_rapidocr,
    "easyocr": _run_easyocr,
    "tesseract": _run_tesseract,
}


def _ocr_pass(original, runner, chosen, preprocess_mode, tier,
             skip_variants=frozenset()):
    """Run every variant for one accuracy tier, skipping names already read."""
    reads = []
    failure = None
    for name, image, opts in build_variants(original, preprocess_mode, tier):
        if name in skip_variants:
            continue
        try:
            segments, took = runner(image, **opts)
        except Exception as exc:
            # One awkward variant should not sink the read, but if they *all*
            # fail the caller must hear about it rather than get empty output.
            failure = exc
            continue
        reads.append(OcrResult(_group_lines(segments), engine=chosen,
                               elapsed=took, variant=name))
    return reads, failure


def read_image(
    data: bytes | str | Path | Image.Image,
    engine: str = "auto",
    preprocess_mode: str = "auto",
    accuracy: str = "fast",
    auto_base: str = "fast",
) -> OcrResult:
    """
    OCR an image and return grouped, ordered text.

    `accuracy` decides how many differently-prepared copies of the image get
    read:

        fast      2 passes  -- cleaned and raw, best one wins. No voting.
        balanced  4 passes  -- plus a big upscale and a loose-threshold read.
        high      7 passes  -- plus grayscale, binarised and a very large scale.
        auto      starts at `auto_base` and works up from there (see below)

    Above `fast` every pass is kept on `.variants`, and extract_fields() votes
    across them. That is where the accuracy comes from: one pass misreading a
    single character is outvoted by the passes that read it correctly.

    With accuracy="auto", the image is read at `auto_base` first, and the
    read's own confidence and word count decide whether that was enough. A
    thin result -- low confidence, very few words -- escalates to the next
    tier and re-reads (only the variants not already tried), repeating until
    the result looks adequate or "high" is reached.

    A blur/resolution pre-check was tried as a way to pick the starting tier
    up front instead of reacting after the fact, and was dropped: Laplacian
    variance (the standard focus measure used here) is scale-dependent in both
    directions -- normalising a small image up to a common reference size
    manufactures blur out of real sharpness through interpolation, and reading
    it at its own small native size manufactures sharpness out of real
    blur through coarse pixel quantisation. Reacting to the OCR engine's own
    reported confidence sidesteps the problem entirely, since that number
    reflects what the engine actually saw rather than a proxy statistic
    computed on pixels beforehand. estimate_quality() is kept and still runs
    on every image -- it is a genuinely useful blur/resolution *description*
    for the interface to show ("this looks blurry"), just not a reliable
    enough number to make the tier decision from.
    """
    engines = available_engines()
    if not engines:
        raise NoEngineAvailable(
            "No OCR engine installed. Install one with:\n"
            "    pip install rapidocr-onnxruntime"
        )

    chosen = engines[0] if engine == "auto" else engine
    if chosen not in engines:
        raise NoEngineAvailable(f"OCR engine '{chosen}' is not available. Found: {engines}")

    runner = _RUNNERS[chosen]
    original = load_image(data)
    quality = estimate_quality(original)     # informational only -- see docstring

    tier = auto_base if accuracy == "auto" else accuracy
    reads, failure = _ocr_pass(original, runner, chosen, preprocess_mode, tier)
    escalated = False

    while accuracy == "auto" and tier != "high":
        best_so_far = (max(reads, key=lambda r: (r.word_count, r.mean_conf))
                       if reads else None)
        thin = (best_so_far is None or best_so_far.word_count == 0
               or best_so_far.mean_conf < _ESCALATE_BELOW_CONF)
        if not thin:
            break
        next_tier = TIER_ORDER[TIER_ORDER.index(tier) + 1]
        already = {r.variant for r in reads}
        more, more_failure = _ocr_pass(original, runner, chosen, preprocess_mode,
                                       next_tier, skip_variants=already)
        tier = next_tier
        if more:
            reads += more
            escalated = True
        elif not failure:
            failure = more_failure

    if not reads:
        if failure is not None:
            raise RuntimeError(f"Every OCR pass failed. Last error: {failure}") from failure
        return OcrResult([], engine=chosen, quality=quality, accuracy_used=tier)

    # The richest read becomes the visible one; ties break on confidence.
    best = max(reads, key=lambda r: (r.word_count, r.mean_conf))
    if len(reads) > 1:
        best.variants = reads
    best.quality = quality
    best.accuracy_used = tier
    best.escalated = escalated
    return best
def read_document(
    data: bytes | str | Path | Image.Image,
    engine: str = "auto",
    preprocess_mode: str = "auto",
    accuracy: str = "fast",
    auto_base: str = "fast",
) -> OcrResult:
    """
    Read a whole document -- an image, or every page of a PDF.

    A PDF page that carries its own embedded text is read directly from that
    text, never rendered to an image or run through OCR at all -- see
    _pdf_native_lines for why. Only a page with no usable text layer (a
    scanned or photographed page) falls back to the render-then-OCR path an
    image always uses. A multi-page PDF can freely mix the two, page by page.

    Pages are concatenated into one result so a report split across sheets is
    extracted as a single document. Variants are merged position-by-position
    (page 1's "big" pass joins page 2's "big" pass) so the voting in
    extract_fields still compares like with like -- that only happens between
    pages that both needed OCR, since a native-text page has nothing to vote
    with.

    `accuracy` and `auto_base` are passed straight through to read_image() for
    any page that needs it -- see there for what "auto" does.
    """
    if not is_pdf(data):
        return read_image(data, engine, preprocess_mode, accuracy, auto_base)

    reads: list[OcrResult] = []
    with _open_pdf(data) as doc:
        for index, page in enumerate(doc):
            if index >= MAX_PDF_PAGES:
                break
            native = _pdf_native_lines(page)
            if native is not None:
                reads.append(OcrResult(lines=native, engine="pdf-text",
                                       accuracy_used="text"))
                continue
            pix = page.get_pixmap(dpi=PDF_RENDER_DPI)
            image = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            reads.append(read_image(image, engine, preprocess_mode, accuracy, auto_base))

    if not reads:
        raise RuntimeError("That PDF has no pages that could be read.")
    if len(reads) == 1:
        return reads[0]

    # Named honestly rather than just "the first page's engine" -- a document
    # that mixes a native page with a scanned one used more than one method,
    # and the interface should be able to say so.
    engines_used = {r.engine for r in reads}
    engine_label = next(iter(engines_used)) if len(engines_used) == 1         else "+".join(sorted(engines_used))

    merged = OcrResult(
        lines=[line for read in reads for line in read.lines],
        engine=engine_label,
        elapsed=sum(read.elapsed for read in reads),
        variant=f"{len(reads)} pages",
        # Report the worst page's score and the highest OCR tier any page
        # needed -- the honest summary of what the document cost, not just
        # what happened to page one. A page read natively contributes neither
        # (it has no blur score and ran no OCR tier at all).
        quality=min((r.quality for r in reads if r.quality),
                    key=lambda q: q.get("laplacian", 1e9), default={}),
        accuracy_used=max((r.accuracy_used for r in reads if r.accuracy_used in TIER_ORDER),
                          key=TIER_ORDER.index, default=""),
        escalated=any(r.escalated for r in reads),
    )

    if all(read.variants for read in reads):
        depth = min(len(read.variants) for read in reads)
        merged.variants = [
            OcrResult(
                lines=[line for read in reads for line in read.variants[i].lines],
                engine=reads[0].engine,
                elapsed=sum(read.variants[i].elapsed for read in reads),
                variant=reads[0].variants[i].variant,
            )
            for i in range(depth)
        ]

    return merged


# =========================================================================== #
#  2. EXTRACTION
# =========================================================================== #
#
# Four strategies are tried per field, strongest first:
#
#   inline     "UTR No: 312345678901"          -- label and value in one box
#   same-line  ["UTR No", "312345678901"]      -- label box, value box beside it
#   next-line  "UTR No" / "312345678901"       -- label box, value box below it
#   scan       a bare 312345678901 anywhere    -- pattern only, no label
#
# Each hit carries a confidence so the UI can show which values to trust. A
# labelled hit always outranks a bare pattern scan, and every hit is scaled by
# the OCR engine's own confidence for the box it came from.

# How much each strategy is trusted before OCR confidence is applied.
METHOD_WEIGHT = {
    "inline": 0.98,
    "same-line": 0.94,
    "next-line": 0.86,
    "scan": 0.55,
}

# Letters that OCR most often swaps for digits, and back.
LETTER_TO_DIGIT = str.maketrans({"O": "0", "o": "0", "I": "1", "l": "1", "|": "1",
                                 "S": "5", "B": "8", "Z": "2", "G": "6", "D": "0"})


@dataclass
class Match:
    key: str
    label: str
    value: str
    raw_value: str = ""
    confidence: float = 0.0
    method: str = "scan"
    context: str = ""
    note: str = ""

    @property
    def repaired(self) -> bool:
        return bool(self.raw_value) and self.raw_value != self.value


@dataclass
class FieldRule:
    key: str
    label: str
    aliases: Sequence[str] = ()
    # Permissive pattern used once a label has been found.
    value_pattern: str = r"[A-Za-z0-9][A-Za-z0-9\-/]{3,40}"
    # Patterns for unlabelled scanning, each with a trust weight. 1.0 is a
    # normal shape match; above 1.0 boosts a self-identifying pattern (one whose
    # prefix makes a false positive near impossible) back up out of "unlabelled
    # guess" territory. Empty disables scanning for this field.
    scan_patterns: Sequence[tuple[str, float]] = ()
    # How many lines below a label to keep looking for its value.
    lookahead: int = 2
    # Short aliases ("To", "From") must own their whole box or they match noise.
    strict_label: bool = False
    # True for free-text name fields (patient_name, payee, payer, referred_by).
    # OCR usually boxes a whole name as one segment, but PDF-native text is
    # split word by word ("MRS." / "LATHA" / "RAMAN" as three segments), and
    # the default same-line search takes the value from the first segment that
    # matches and stops there -- fine for a UTR or an amount, which are always
    # one token, but it silently truncated a real name to just its first word.
    # Scoped to exactly the fields that need it so a UTR's tighter matching
    # (and its 30-char cap, meant to stop it running on into unrelated text)
    # stays exactly as tested.
    multi_segment_value: bool = False
    normalise: Callable[[str], str] | None = None
    validate: Callable[[str], bool] | None = None
    # Keep every distinct hit rather than only the strongest one.
    multi: bool = False
    # Lines containing these words are skipped during unlabelled scanning.
    scan_blocklist: Sequence[str] = ()


def _clean_ref(value: str) -> str:
    """Reference numbers: drop spacing/punctuation noise, upper-case."""
    return re.sub(r"[\s\-_.]", "", value).upper()


def _clean_txn(value: str) -> str:
    """
    Payment-gateway IDs are not bank references and must not be normalised
    like one. Razorpay/Stripe ids such as `pay_TR6FFa1mQcOkwM` are
    case-sensitive, and the `pay_` prefix is part of the id -- so only
    whitespace comes out, and the case is left exactly as read.
    """
    return re.sub(r"\s+", "", value).strip(" :-#.,")


def _clean_amount(value: str) -> str:
    """
    Normalise a money string to plain digits and one decimal point.

    OCR routinely reads the thousands comma in "3,999.00" as a period, so the
    two separators cannot be told apart by character. Position settles it: the
    last separator is the decimal point only when exactly two digits follow it,
    and every other separator is a group separator. Getting this wrong turns
    3,999.00 into 3.99, which is a silently wrong number rather than a visibly
    missing one.
    """
    token = re.sub(r"[^\d.,]", "", value)
    if not token:
        return ""

    decimal = re.search(r"[.,](\d{2})$", token)
    if decimal:
        whole = re.sub(r"\D", "", token[:decimal.start()])
        return f"{whole}.{decimal.group(1)}" if whole else f"0.{decimal.group(1)}"
    return re.sub(r"\D", "", token)


def _clean_plain(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip(" :-#.,")


def _valid_utr(value: str) -> bool:
    """
    Accept the shapes Indian payment rails actually issue:
      12 digits            UPI / IMPS
      16 alphanumeric      NEFT   (4-letter bank code + 12)
      22 alphanumeric      RTGS
      9-24 alphanumeric    other bank references, as a catch-all
    """
    v = _clean_ref(value)
    if not (9 <= len(v) <= 24):
        return False
    if not any(ch.isdigit() for ch in v):
        return False
    return bool(re.fullmatch(r"[A-Z0-9]+", v))


def classify_utr(value: str) -> str:
    """Best guess at which rail issued a reference, for display only."""
    v = _clean_ref(value)
    if re.fullmatch(r"\d{12}", v):
        return "UPI / IMPS (12-digit)"
    if re.fullmatch(r"[A-Z]{4}[A-Z0-9]{12}", v):
        return "NEFT (16-char)"
    if re.fullmatch(r"[A-Z]{4}[A-Z0-9]{18}", v):
        return "RTGS (22-char)"
    if re.fullmatch(r"\d{16}", v):
        return "NEFT (16-digit)"
    return "Bank reference"


DATE_PATTERN = (
    r"(?:\d{1,2}[/\-.]\d{1,2}[/\-.]\d{2,4}"
    r"|\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?,?\s+\d{2,4}"
    r"|(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{1,2},?\s+\d{2,4}"
    r"|\d{4}-\d{2}-\d{2})"
)
TIME_PATTERN = r"\d{1,2}:\d{2}(?::\d{2})?\s*(?:[AaPp]\.?[Mm]\.?)?"

BUILTIN_FIELDS: list[FieldRule] = [
    FieldRule(
        key="utr",
        label="UTR / Reference Number",
        aliases=["utr", "utr no", "utr number", "utr ref", "unique transaction reference",
                 "bank reference", "bank ref no", "reference no", "reference number",
                 "ref no", "ref id", "rrn", "retrieval reference"],
        value_pattern=r"[A-Za-z0-9][A-Za-z0-9\s\-]{7,30}",
        scan_patterns=[
            (r"\b(?:[A-Z]{4}[A-Z0-9]{18}|[A-Z]{4}[A-Z0-9]{12}|\d{16}|\d{12})\b", 1.0),
        ],
        normalise=_clean_ref,
        validate=_valid_utr,
        multi=True,
        scan_blocklist=["mobile", "phone", "card", "aadhaar", "aadhar", "pan", "gst"],
    ),
    FieldRule(
        key="txn_id",
        label="Transaction ID",
        aliases=["transaction id", "transaction no", "transaction number", "txn id",
                 "txn no", "upi transaction id", "upi ref no", "upi reference id",
                 "order id", "payment id", "receipt no", "invoice no"],
        value_pattern=r"[A-Za-z0-9][A-Za-z0-9\s\-_/]{5,40}",
        # Gateway ids carry their own prefix, so they are safe to spot without a
        # label -- on a Razorpay receipt the id often sits under "UPI" or no
        # label at all, which no alias would ever catch.
        scan_patterns=[
            # Boosted: the prefix makes this all but impossible to hit by
            # accident, so it deserves more trust than a bare shape match.
            (r"\b(?:pay|order|rfnd|plink|inv|txn|pi|ch|cus|sub)_[A-Za-z0-9]{8,24}\b",
             1.6),
        ],
        normalise=_clean_txn,
        multi=True,
    ),
    FieldRule(
        key="amount",
        label="Amount",
        # "paid" alone is deliberately absent -- it swallows "Paid To <name>".
        aliases=["amount paid", "paid amount", "amount", "amt", "grand total",
                 "total amount", "total", "transaction amount", "debited",
                 "credited", "sum"],
        # Separators are matched as a class rather than spelled out, because OCR
        # swaps "," and "." freely; _clean_amount decides which was the decimal.
        value_pattern=r"(?:INR|Rs\.?|₹)?\s*\d(?:[\d.,]*\d)?",
        # OCR frequently mangles the rupee glyph (it comes back as 7, 3, or a CJK
        # character), so money-shaped numbers are matched on their own too --
        # just trusted a little less than an explicit currency prefix.
        scan_patterns=[
            (r"(?:₹|INR|Rs\.?)\s*\d(?:[\d.,]*\d)?", 1.0),
            # A stray symbol or non-ASCII glyph in front of a money-shaped number
            # is almost always a rupee sign the OCR failed to name.
            (r"(?:^|\s)(?:[^\w\s]|[^\x00-\x7F])\s*\d{1,3}(?:[.,]\d{2,3})+", 0.95),
            (r"\b\d{1,3}(?:[.,]\d{2,3})+\b", 0.85),   # 4,250.00 or 4.250,00
            (r"\b\d+[.,]\d{2}\b", 0.7),               # 4250.00
        ],
        lookahead=1,        # amounts sit next to their label, never two rows down
        normalise=_clean_amount,
        multi=True,
    ),
    FieldRule(
        key="date",
        label="Date",
        aliases=["date", "date & time", "date and time", "transaction date",
                 "paid on", "on", "timestamp", "value date"],
        value_pattern=DATE_PATTERN,
        scan_patterns=[(DATE_PATTERN, 1.0)],
        normalise=_clean_plain,
    ),
    FieldRule(
        key="time",
        label="Time",
        aliases=["time", "transaction time"],
        value_pattern=TIME_PATTERN,
        scan_patterns=[(TIME_PATTERN, 1.0)],
        normalise=_clean_plain,
    ),
    FieldRule(
        key="ifsc",
        label="IFSC Code",
        aliases=["ifsc", "ifsc code", "ifs code"],
        value_pattern=r"[A-Za-z]{4}0[A-Za-z0-9]{6}",
        scan_patterns=[(r"\b[A-Z]{4}0[A-Z0-9]{6}\b", 1.0)],
        normalise=_clean_ref,
        multi=True,
    ),
    FieldRule(
        key="upi_id",
        label="UPI ID (VPA)",
        aliases=["upi id", "vpa", "virtual payment address", "upi handle"],
        value_pattern=r"[A-Za-z0-9._\-]{2,}@[A-Za-z]{2,}",
        scan_patterns=[(r"\b[A-Za-z0-9._\-]{2,}@[A-Za-z]{2,}\b(?!\.[A-Za-z])", 1.0)],
        normalise=lambda v: v.strip().lower(),
        multi=True,
    ),
    FieldRule(
        key="account",
        label="Account Number",
        aliases=["account", "account no", "account number", "a/c", "a/c no",
                 "acct no", "bank account", "debited from", "credited to"],
        value_pattern=r"(?:[Xx*•]{2,}\s?\d{3,6}|\d{6,18})",
        scan_patterns=[(r"\b(?:[Xx*]{4,}\d{3,6})\b", 1.0)],
        normalise=lambda v: re.sub(r"\s", "", v).upper(),
        multi=True,
    ),
    FieldRule(
        key="payee",
        label="Paid To / Beneficiary",
        aliases=["paid to", "to", "payee", "beneficiary", "receiver", "recipient",
                 "transferred to", "sent to", "merchant"],
        value_pattern=r"[A-Za-z0-9][\w\s.&'@\-]{2,60}",
        strict_label=True, multi_segment_value=True,
        normalise=_clean_plain,
    ),
    FieldRule(
        key="payer",
        label="Paid By / Sender",
        aliases=["paid by", "from", "payer", "sender", "debited from account",
                 "remitter", "sent from"],
        value_pattern=r"[A-Za-z0-9][\w\s.&'@\-]{2,60}",
        strict_label=True, multi_segment_value=True,
        normalise=_clean_plain,
    ),
    FieldRule(
        key="status",
        label="Status",
        aliases=["status", "transaction status", "payment status"],
        value_pattern=r"[A-Za-z][A-Za-z\s]{2,25}",
        scan_patterns=[
            (r"\b(?:Success(?:ful)?|Completed|Complete|Paid|Failed|Failure"
             r"|Pending|Processing|Declined|Rejected|Cancelled)\b", 1.0),
        ],
        normalise=_clean_plain,
    ),
    FieldRule(
        key="mode",
        label="Payment Mode",
        aliases=["mode", "payment mode", "transfer type", "transaction type", "via"],
        value_pattern=r"[A-Za-z][A-Za-z\s/]{1,25}",
        scan_patterns=[
            (r"\b(?:UPI|IMPS|NEFT|RTGS|Net\s?Banking|Debit\s?Card"
             r"|Credit\s?Card|Wallet)\b", 1.0),
        ],
        normalise=_clean_plain,
    ),
]

FIELDS_BY_KEY = {rule.key: rule for rule in BUILTIN_FIELDS}
DEFAULT_ENABLED = ["utr", "txn_id", "amount", "date", "payee", "payer", "status"]


def _label_regex(aliases: Iterable[str], strict: bool) -> re.Pattern:
    """
    Build a regex that matches a label at the start of an OCR box.

    Aliases are sorted longest-first so "utr number" wins over "utr", and the
    trailing group soaks up the "No. :" style punctuation that sits between a
    label and its value.
    """
    parts = sorted((re.escape(a).replace(r"\ ", r"\s*") for a in aliases),
                   key=len, reverse=True)
    body = r"|".join(parts)
    tail = r"\s*(?:number|no\.?|id|#)?\s*[:\-–—#>.]*\s*"
    # A stray leading bracket or dash shows up on real reports -- confirmed on
    # one where OCR read a label box as ") Rpt Date : ...". Tolerated only from
    # a small explicit set, and only up to 3 characters, so this stays a tight
    # anchor rather than turning into an unanchored search that could match a
    # label alias appearing mid-sentence.
    head = r"^[\s)(\]\[.,:;\-]{0,3}"
    if strict:
        # Whole box must be the label (plus punctuation) -- nothing else.
        return re.compile(rf"{head}(?:{body}){tail}$", re.IGNORECASE)
    return re.compile(rf"{head}(?:{body})\b{tail}", re.IGNORECASE)


def _repair_digits(value: str, rule: FieldRule) -> tuple[str, str]:
    """
    If a value is *nearly* all digits, assume the stray letters are OCR slips
    and map them back (O->0, I->1, S->5...). Returns (repaired, note).
    """
    if rule.key not in ("utr", "txn_id", "account"):
        return value, ""

    stripped = _clean_ref(value)
    letters = sum(ch.isalpha() for ch in stripped)
    digits = sum(ch.isdigit() for ch in stripped)
    # Only worth repairing when digits clearly dominate.
    if digits >= 8 and 0 < letters <= 2:
        repaired = stripped.translate(LETTER_TO_DIGIT)
        if repaired != stripped and repaired.isdigit():
            return repaired, f"OCR repair applied ({stripped} -> {repaired})"
    return value, ""


def _build_match(rule: FieldRule, raw: str, method: str, line: Line,
                 ocr_conf: float, decay: float = 1.0) -> Match | None:
    value = raw.strip()
    if not value:
        return None

    if rule.normalise:
        value = rule.normalise(value)
    if not value:
        return None

    value, note = _repair_digits(value, rule)

    if rule.validate and not rule.validate(value):
        return None

    confidence = METHOD_WEIGHT.get(method, 0.5) * max(ocr_conf, 0.35) * decay
    if note:
        confidence *= 0.95      # a repaired value is slightly less certain

    return Match(
        key=rule.key,
        label=rule.label,
        value=value,
        raw_value=raw.strip(),
        confidence=round(min(confidence, 0.99), 3),
        method=method,
        context=line.text[:120],
        note=note,
    )


def _extract_labeled(rule: FieldRule, lines: Sequence[Line]) -> list[Match]:
    label_re = _label_regex(rule.aliases, rule.strict_label)
    value_re = re.compile(rule.value_pattern, re.IGNORECASE)
    found: list[Match] = []

    for i, line in enumerate(lines):
        for j, seg in enumerate(line.segments):
            hit = label_re.match(seg.text)
            if not hit:
                continue

            # 1. value trails the label inside the same box
            remainder = seg.text[hit.end():].strip()
            if remainder:
                vm = value_re.search(remainder)
                if vm:
                    m = _build_match(rule, vm.group(0), "inline", line, seg.conf)
                    if m:
                        found.append(m)
                        continue

            # 2. value sits in a later box on the same visual line
            placed = False
            later_segments = line.segments[j + 1:]
            if rule.multi_segment_value:
                # A name can be split across several word-level segments (PDF-
                # native text does this; OCR usually does not). Join them and
                # match once, rather than stopping at whichever segment
                # happens to satisfy the pattern first and silently dropping
                # the rest of the name.
                joined = " ".join(s.text for s in later_segments).strip()
                if joined:
                    vm = value_re.search(joined)
                    if vm:
                        m = _build_match(rule, vm.group(0), "same-line", line,
                                         line.conf)
                        if m:
                            found.append(m)
                            placed = True
            else:
                for later in later_segments:
                    vm = value_re.search(later.text)
                    if vm:
                        m = _build_match(rule, vm.group(0), "same-line", line,
                                         later.conf)
                        if m:
                            found.append(m)
                            placed = True
                            break
            if placed:
                continue

            # 3. value sits on one of the next lines
            for step, k in enumerate(range(i + 1, min(i + 1 + rule.lookahead, len(lines)))):
                nxt = lines[k]
                if not nxt.segments:
                    continue
                vm = value_re.search(nxt.text)
                if vm:
                    m = _build_match(rule, vm.group(0), "next-line", nxt,
                                     nxt.conf, decay=1.0 - 0.15 * step)
                    if m:
                        found.append(m)
                        break

    return found


def _extract_scan(rule: FieldRule, lines: Sequence[Line]) -> list[Match]:
    """Find values by shape alone, with no label to anchor them."""
    if not rule.scan_patterns:
        return []

    compiled = [(re.compile(pattern), weight) for pattern, weight in rule.scan_patterns]
    found: list[Match] = []

    for line in lines:
        lowered = line.text.lower()
        if any(word in lowered for word in rule.scan_blocklist):
            continue

        # Patterns run strongest-first, so once a stretch of the line is claimed
        # a weaker pattern must not re-report part of it (".00" out of "4,250.00").
        claimed: list[tuple[int, int]] = []

        for scan_re, weight in compiled:
            for vm in scan_re.finditer(line.text):
                start, end = vm.span()
                if any(start < c_end and c_start < end for c_start, c_end in claimed):
                    continue
                m = _build_match(rule, vm.group(0), "scan", line, line.conf,
                                 decay=weight)
                if not m:
                    continue
                claimed.append((start, end))
                # A bare 12-digit token shaped like +91 mobile is probably not a UTR.
                if rule.key == "utr" and re.fullmatch(r"91[6-9]\d{9}", m.value):
                    m.confidence = round(m.confidence * 0.4, 3)
                    m.note = "Looks like a phone number -- verify before use"
                found.append(m)

    return found


def _dedupe(matches: Iterable[Match]) -> list[Match]:
    """Keep the highest-confidence hit per distinct value."""
    best: dict[str, Match] = {}
    for m in matches:
        key = m.value.upper()
        if key not in best or m.confidence > best[key].confidence:
            best[key] = m
    return sorted(best.values(), key=lambda m: m.confidence, reverse=True)


# Fields whose values are fixed-shape identifiers, where voting one character
# at a time is meaningful. Names, dates and amounts are excluded -- their length
# legitimately varies, so aligning them position-by-position is nonsense.
CHAR_VOTE_FIELDS = {"utr", "txn_id", "account", "ifsc"}

# Glyphs that render near-identically, so OCR picks between them by guessing.
# Measured on this engine, every ID miss was one of these -- 0 read as O or o,
# lowercase l read as uppercase I. Crucially the mistake is *systematic*: every
# preprocessing variant makes the same call, so no amount of voting fixes it.
# The only honest answer is to surface the other reading as a candidate.
# Tier 1: digit/letter pairs. Every miss the benchmark recorded was one of
# these, and the replacements are ordered most-likely-first -- when the engine
# prints "o" inside an id the intended character is far more often a zero than a
# capital O, so "0" is offered before "O".
_TIER1: dict[str, tuple[str, ...]] = {
    "0": ("O", "o"), "O": ("0", "o"), "o": ("0", "O"), "D": ("0",),
    "1": ("l", "I"), "l": ("1", "I"), "I": ("l", "1"), "|": ("1", "l"),
}
# Tier 2: same-shape case pairs and rarer digit/letter swaps. Real but much less
# frequent, so they are only reached once tier 1 is exhausted.
_TIER2_GROUPS = ("5S", "2Z", "8B", "6G", "9gq", "7T",
                 "cC", "kK", "pP", "sS", "uU", "vV", "wW", "xX", "yY", "zZ")
_TIER2: dict[str, list[str]] = {}
for _group in _TIER2_GROUPS:
    for _ch in _group:
        _TIER2.setdefault(_ch, []).extend(c for c in _group if c != _ch)

_CONFUSABLE: dict[str, list[str]] = {
    ch: list(subs) for ch, subs in _TIER1.items()
}
for _ch, _subs in _TIER2.items():
    _CONFUSABLE.setdefault(_ch, []).extend(_subs)


def ambiguous_positions(value: str) -> list[int]:
    """Indexes holding a character this engine is known to guess at."""
    return [i for i, ch in enumerate(value) if ch in _CONFUSABLE]


def _worth_disambiguating(m: Match) -> bool:
    """
    Whether offering look-alike readings of a value helps or just adds noise.

    An all-digit id has a known charset, so a stray letter in it is a repair
    rather than a genuine choice -- offering "3l2845967103" next to a confident
    "312845967103" would be pure clutter. The case that genuinely needs the
    user's eye is a mixed-case value, where 0/O/o are indistinguishable and
    nothing in the format can settle it.
    """
    value = m.value
    if value.isdigit():
        return False
    if any(c.islower() for c in value) and any(c.isupper() for c in value):
        return True
    return m.confidence < 0.85


def ambiguity_alternates(value: str, limit: int = 7) -> list[str]:
    """
    The same ID with its ambiguous glyphs swapped for their look-alikes.

    Ordered by how often each substitution actually turned out to be the fix:

      1. single swaps at tier-1 (digit/letter) positions
      2. *pairs* of tier-1 swaps -- a measured miss needed both o->0 and I->l,
         so single swaps alone could never reach the right string
      3. single swaps at tier-2 (case-pair) positions

    A literal gateway prefix is skipped entirely. "pay_" is fixed text, so
    offering "Pay_" and "paY_" only burns the budget before the generator ever
    reaches the characters that are genuinely in doubt.
    """
    start = value.find("_") + 1        # 0 when there is no prefix
    body = range(start, len(value))
    tier1 = [i for i in body if value[i] in _TIER1]
    tier2 = [i for i in body if value[i] in _TIER2 and value[i] not in _TIER1]

    out: list[str] = []

    def offer(candidate: str) -> bool:
        if candidate != value and candidate not in out:
            out.append(candidate)
        return len(out) >= limit

    for i in tier1:
        for swap in _TIER1[value[i]]:
            if offer(value[:i] + swap + value[i + 1:]):
                return out

    for a in range(len(tier1)):
        for b in range(a + 1, len(tier1)):
            i, j = tier1[a], tier1[b]
            for si in _TIER1[value[i]]:
                for sj in _TIER1[value[j]]:
                    chars = list(value)
                    chars[i], chars[j] = si, sj
                    if offer("".join(chars)):
                        return out

    for i in tier2:
        for swap in _TIER2[value[i]]:
            if offer(value[:i] + swap + value[i + 1:]):
                return out

    return out


def _extract_single(rules: Sequence[FieldRule], ocr: OcrResult) -> dict[str, list[Match]]:
    """Every rule against one OCR read."""
    results: dict[str, list[Match]] = {}
    for rule in rules:
        hits = _extract_labeled(rule, ocr.lines)
        # Only fall back to blind scanning when the label search came up empty.
        if not hits or rule.multi:
            hits.extend(_extract_scan(rule, ocr.lines))

        merged = _dedupe(hits)
        if merged:
            results[rule.key] = merged
    return results


def _char_vote(rule: FieldRule, candidates: Sequence[Match],
               support: Counter) -> tuple[str, float] | None:
    """
    Build a consensus string by taking a weighted majority at each position.

    This is what recovers an ID no single pass read correctly: if one pass says
    ...96I103 and three say ...961103, the majority wins that position while
    every other position stays untouched. Only same-length candidates are
    aligned, since a dropped character would shift everything after it.
    """
    if len(candidates) < 3:
        return None

    lengths: Counter = Counter()
    for m in candidates:
        lengths[len(m.value)] += support.get(m.value, 1)
    length = lengths.most_common(1)[0][0]

    pool = [m for m in candidates if len(m.value) == length]
    if len(pool) < 3:
        return None

    chars, agreements = [], []
    for i in range(length):
        tally: dict[str, float] = defaultdict(float)
        for m in pool:
            tally[m.value[i]] += max(m.confidence, 0.05) * support.get(m.value, 1)
        char, weight = max(tally.items(), key=lambda kv: kv[1])
        total = sum(tally.values())
        chars.append(char)
        agreements.append(weight / total if total else 0.0)

    voted = "".join(chars)
    if rule.validate and not rule.validate(voted):
        return None

    # Confidence is the mean per-position agreement: a string where every
    # position was unanimous scores far above one stitched from close calls.
    return voted, round(min(0.97, (sum(agreements) / len(agreements)) * 0.95), 3)


def _consensus(rules: Sequence[FieldRule],
               per_variant: Sequence[dict[str, list[Match]]]) -> dict[str, list[Match]]:
    """Merge several reads of the same image into one, agreement-weighted."""
    n = len(per_variant)
    out: dict[str, list[Match]] = {}

    for rule in rules:
        support: Counter = Counter()
        best_by_value: dict[str, Match] = {}

        for read in per_variant:
            hits = read.get(rule.key, [])
            for value in {m.value for m in hits}:      # one vote per read
                support[value] += 1
            for m in hits:
                current = best_by_value.get(m.value)
                if current is None or m.confidence > current.confidence:
                    best_by_value[m.value] = m

        if not best_by_value:
            continue

        scored: list[Match] = []
        for value, m in best_by_value.items():
            agree = support[value] / n
            # Unanimity earns a modest boost; a lone dissenting read is damped.
            confidence = round(min(0.99, m.confidence * (0.55 + 0.55 * agree)), 3)
            note = m.note
            if n >= 3 and support[value] == 1 and not note:
                note = f"only 1 of {n} reads found this"
            scored.append(replace(m, confidence=confidence, note=note))

        if rule.key in CHAR_VOTE_FIELDS:
            voted = _char_vote(rule, list(best_by_value.values()), support)
            if voted:
                value, confidence = voted
                if value in best_by_value:
                    # Voting agreed with a value we already have: that is
                    # corroboration, so let it keep the higher score.
                    for i, m in enumerate(scored):
                        if m.value == value:
                            scored[i] = replace(m, confidence=max(m.confidence, confidence))
                            break
                else:
                    proto = max(best_by_value.values(), key=lambda m: m.confidence)
                    scored.append(Match(
                        key=rule.key, label=rule.label, value=value,
                        raw_value=proto.raw_value, confidence=confidence,
                        method="consensus", context=proto.context,
                        note=f"assembled from {n} reads that disagreed"))

        scored.sort(key=lambda m: m.confidence, reverse=True)

        out[rule.key] = scored

    return out


def _add_lookalikes(rules: Sequence[FieldRule],
                    found: dict[str, list[Match]]) -> dict[str, list[Match]]:
    """
    Append alternative readings of ambiguous glyphs to ID fields.

    Applied after merging rather than inside the vote, so it works the same
    whether one read happened or seven -- the ambiguity is a property of the
    glyphs, not of how many times they were read.
    """
    for rule in rules:
        scored = found.get(rule.key)
        if rule.key not in CHAR_VOTE_FIELDS or not scored:
            continue
        top = scored[0]
        if not _worth_disambiguating(top):
            continue

        known = {m.value for m in scored}
        swapped: set[str] = set()

        for alt in ambiguity_alternates(top.value):
            if alt in known:
                continue
            if rule.validate and not rule.validate(alt):
                continue
            swapped.update(a for a, b in zip(top.value, alt) if a != b)
            scored.append(replace(
                top, value=alt, confidence=round(top.confidence * 0.45, 3),
                method="look-alike",
                note="look-alike glyphs swapped -- click if this is the real id"))

        if swapped and not top.note:
            scored[0] = replace(
                top, note=f"{', '.join(sorted(swapped))} could be misread here "
                          f"-- check before use")

    return found


def extract_fields(
    ocr: OcrResult,
    enabled: Sequence[str] = tuple(DEFAULT_ENABLED),
    custom_rules: Sequence[FieldRule] = (),
) -> dict[str, list[Match]]:
    """
    Run every enabled rule over an OCR result.

    When the read carries several variants (accuracy above "fast"), each is
    extracted separately and the results are voted on. Returns
    {field_key: [Match, ...]} ordered by confidence; single-value fields are
    trimmed to their best hit, multi-value fields keep all distinct candidates.
    """
    rules = [FIELDS_BY_KEY[k] for k in enabled if k in FIELDS_BY_KEY]
    rules.extend(custom_rules)

    reads = ocr.variants or [ocr]
    per_variant = [_extract_single(rules, r) for r in reads]
    merged = per_variant[0] if len(per_variant) == 1 else _consensus(rules, per_variant)
    merged = _add_lookalikes(rules, merged)

    return {rule.key: (merged[rule.key] if rule.multi else merged[rule.key][:1])
            for rule in rules if merged.get(rule.key)}


class UnsafePattern(ValueError):
    """A user-supplied regex that could hang the process."""


# Nested quantifiers -- (a+)+ , (a*)* , (\d+|x)+ -- are the classic recipe for
# catastrophic backtracking. Python's re has no timeout and cannot be
# interrupted, so one bad pattern typed into the sidebar would wedge the OCR
# thread for good. Refusing the shape is cheaper than trying to survive it.
_NESTED_QUANTIFIER = re.compile(r"\([^)]*[+*][^)]*\)\s*[+*{]")


def safe_compile(pattern: str, limit: int = 200) -> re.Pattern:
    """Compile a user-supplied pattern, refusing the ones known to blow up."""
    if len(pattern) > limit:
        raise UnsafePattern(
            f"Pattern is too long ({len(pattern)} characters, limit {limit}).")
    if _NESTED_QUANTIFIER.search(pattern):
        raise UnsafePattern(
            "Pattern nests one repeat inside another, like (a+)+. That can take "
            "effectively forever to match, so it is refused. Rewrite it without "
            "the inner repeat.")
    try:
        return re.compile(pattern)
    except re.error as exc:
        raise UnsafePattern(f"Not a valid regular expression: {exc}") from exc


def make_custom_rule(label: str, keywords: str = "", pattern: str = "") -> FieldRule:
    """
    Build a rule from sidebar input.

    Either half is optional: keywords alone searches by label, a pattern alone
    scans the whole image, and both together does label-first then scan.
    """
    pattern = pattern.strip()
    if pattern:
        safe_compile(pattern)      # raises UnsafePattern, caught by the caller

    key = re.sub(r"[^a-z0-9]+", "_", label.strip().lower()).strip("_") or "custom"
    aliases = [a.strip() for a in re.split(r"[,;|]", keywords) if a.strip()]
    if not aliases:
        aliases = [label.strip()]

    return FieldRule(
        key=key,
        label=label.strip() or "Custom Field",
        aliases=aliases,
        value_pattern=pattern or r"[A-Za-z0-9][\w\s.\-/@]{2,60}",
        # A user-supplied pattern is specific enough to scan for on its own;
        # the generic fallback pattern is not, so it stays label-only.
        scan_patterns=[(pattern, 1.0)] if pattern else (),
        normalise=_clean_plain,
        multi=True,
    )


# =========================================================================== #
#  3. EXPORT
# =========================================================================== #

@dataclass
class ExtractionRecord:
    """Everything extracted from a single image."""
    filename: str
    fields: dict[str, list[Match]]
    full_text: str = ""
    engine: str = ""
    ocr_confidence: float = 0.0
    elapsed: float = 0.0

    def primary(self, key: str) -> str:
        """Best value for a field, or empty string."""
        hits = self.fields.get(key)
        return hits[0].value if hits else ""

    def flat_rows(self) -> list[dict]:
        """One row per extracted value -- the shape CSV and tables want."""
        rows = []
        for hits in self.fields.values():
            for rank, m in enumerate(hits, start=1):
                rows.append({
                    "file": self.filename,
                    "field": m.label,
                    "value": m.value,
                    "confidence": round(m.confidence, 3),
                    "rank": rank,
                    "method": m.method,
                    "raw_value": m.raw_value,
                    "note": m.note,
                    "context": m.context,
                })
        return rows


def _timestamp() -> str:
    return datetime.now().strftime("%d %b %Y, %I:%M %p")


def _confidence_word(conf: float) -> str:
    if conf >= 0.8:
        return "High"
    if conf >= 0.6:
        return "Medium"
    return "Low"


def _latin1_safe(text: str) -> str:
    """Strip glyphs the default PDF fonts cannot draw (rupee sign, dashes)."""
    replacements = {"₹": "INR ", "–": "-", "—": "-",
                    "‘": "'", "’": "'", "“": '"', "”": '"',
                    "•": "-"}
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.encode("latin-1", "replace").decode("latin-1")


def build_txt(records: Sequence[ExtractionRecord], include_raw_text: bool = False) -> bytes:
    out = io.StringIO()
    out.write(f"{APP_TITLE}\n")
    out.write("=" * len(APP_TITLE) + "\n")
    out.write(f"Generated : {_timestamp()}\n")
    out.write(f"Images    : {len(records)}\n\n")

    for idx, rec in enumerate(records, start=1):
        out.write(f"\n[{idx}] {rec.filename}\n")
        out.write("-" * (len(rec.filename) + 6) + "\n")
        out.write(f"OCR engine: {rec.engine}   confidence: {rec.ocr_confidence:.0%}\n\n")

        if not rec.fields:
            out.write("  No fields matched.\n")
        for hits in rec.fields.values():
            for m in hits:
                out.write(f"  {m.label:<26}: {m.value}\n")
                detail = f"({_confidence_word(m.confidence)} {m.confidence:.0%}, via {m.method})"
                out.write(f"  {'':<26}  {detail}\n")
                if m.note:
                    out.write(f"  {'':<26}  ! {m.note}\n")

        if include_raw_text and rec.full_text:
            out.write("\n  --- Full OCR text ---\n")
            for line in rec.full_text.splitlines():
                out.write(f"  {line}\n")

    return out.getvalue().encode("utf-8")


def build_csv(records: Sequence[ExtractionRecord]) -> bytes:
    columns = ["file", "field", "value", "confidence", "rank", "method",
               "raw_value", "note", "context"]
    out = io.StringIO(newline="")
    writer = csv.DictWriter(out, fieldnames=columns, extrasaction="ignore")
    writer.writeheader()
    for rec in records:
        writer.writerows(rec.flat_rows())
    # BOM so Excel opens UTF-8 cleanly on Windows.
    return out.getvalue().encode("utf-8-sig")


def build_json(records: Sequence[ExtractionRecord], include_raw_text: bool = True) -> bytes:
    payload = {
        "tool": APP_TITLE,
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "image_count": len(records),
        "results": [],
    }

    for rec in records:
        entry = {
            "file": rec.filename,
            "ocr_engine": rec.engine,
            "ocr_confidence": round(rec.ocr_confidence, 3),
            "seconds": round(rec.elapsed, 3),
            "fields": {
                key: [
                    {
                        "value": m.value,
                        "confidence": m.confidence,
                        "method": m.method,
                        "raw_value": m.raw_value,
                        "note": m.note,
                    }
                    for m in hits
                ]
                for key, hits in rec.fields.items()
            },
        }
        if include_raw_text:
            entry["full_text"] = rec.full_text
        payload["results"].append(entry)

    return json.dumps(payload, indent=2, ensure_ascii=False).encode("utf-8")


def build_docx(records: Sequence[ExtractionRecord], include_raw_text: bool = False) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    doc.add_heading(APP_TITLE, level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {_timestamp()}\n").italic = True
    meta.add_run(f"Images processed: {len(records)}").italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for idx, rec in enumerate(records, start=1):
        doc.add_heading(f"{idx}. {rec.filename}", level=1)

        sub = doc.add_paragraph()
        run = sub.add_run(
            f"OCR engine: {rec.engine}    "
            f"OCR confidence: {rec.ocr_confidence:.0%}    "
            f"Time: {rec.elapsed:.2f}s"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        if not rec.fields:
            doc.add_paragraph("No fields matched in this image.")
            continue

        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        for cell, heading in zip(table.rows[0].cells,
                                 ["Field", "Value", "Confidence", "Found via"]):
            cell.text = heading
            for para in cell.paragraphs:
                for r in para.runs:
                    r.bold = True

        for hits in rec.fields.values():
            for rank, m in enumerate(hits):
                cells = table.add_row().cells
                cells[0].text = m.label if rank == 0 else f"{m.label} (alt {rank + 1})"
                cells[1].text = m.value
                cells[2].text = f"{_confidence_word(m.confidence)} ({m.confidence:.0%})"
                cells[3].text = m.method

                # Value column is the one people copy -- make it stand out.
                for para in cells[1].paragraphs:
                    for r in para.runs:
                        r.bold = True

        notes = [m for hits in rec.fields.values() for m in hits if m.note]
        if notes:
            doc.add_paragraph()
            for m in notes:
                note_para = doc.add_paragraph(f"Note on {m.label}: {m.note}",
                                              style="List Bullet")
                for r in note_para.runs:
                    r.font.size = Pt(9)

        if include_raw_text and rec.full_text:
            doc.add_heading("Full OCR text", level=2)
            block = doc.add_paragraph(rec.full_text)
            for r in block.runs:
                r.font.name = "Consolas"
                r.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_pdf(records: Sequence[ExtractionRecord], include_raw_text: bool = False) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.platypus import (PageBreak, Paragraph, SimpleDocTemplate,
                                    Spacer, Table, TableStyle)

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=18 * mm, rightMargin=18 * mm,
        topMargin=16 * mm, bottomMargin=16 * mm,
        title=APP_TITLE,
    )

    styles = getSampleStyleSheet()
    style_meta = ParagraphStyle("meta", parent=styles["Normal"],
                                fontSize=8.5, textColor=colors.HexColor("#666666"))
    style_cell = ParagraphStyle("cell", parent=styles["Normal"], fontSize=9, leading=12)
    style_value = ParagraphStyle("value", parent=style_cell, fontName="Helvetica-Bold")
    style_mono = ParagraphStyle("mono", parent=styles["Normal"],
                                fontName="Courier", fontSize=7.5, leading=9.5)

    story = [
        Paragraph(APP_TITLE, styles["Title"]),
        Paragraph(f"Generated: {_timestamp()} &nbsp;|&nbsp; Images processed: {len(records)}",
                  style_meta),
        Spacer(1, 8 * mm),
    ]

    for idx, rec in enumerate(records, start=1):
        story.append(Paragraph(f"{idx}. {escape(_latin1_safe(rec.filename))}",
                               styles["Heading2"]))
        story.append(Paragraph(
            f"OCR engine: {rec.engine} &nbsp;|&nbsp; "
            f"OCR confidence: {rec.ocr_confidence:.0%} &nbsp;|&nbsp; "
            f"Time: {rec.elapsed:.2f}s",
            style_meta,
        ))
        story.append(Spacer(1, 3 * mm))

        if not rec.fields:
            story.append(Paragraph("No fields matched in this image.", style_cell))
            story.append(Spacer(1, 6 * mm))
            continue

        data = [[Paragraph(f"<b>{h}</b>", style_cell)
                 for h in ("Field", "Value", "Confidence", "Found via")]]

        for hits in rec.fields.values():
            for rank, m in enumerate(hits):
                name = m.label if rank == 0 else f"{m.label} (alt {rank + 1})"
                data.append([
                    Paragraph(escape(_latin1_safe(name)), style_cell),
                    Paragraph(escape(_latin1_safe(m.value)), style_value),
                    Paragraph(f"{_confidence_word(m.confidence)} ({m.confidence:.0%})",
                              style_cell),
                    Paragraph(m.method, style_cell),
                ])

        table = Table(data, colWidths=[46 * mm, 68 * mm, 30 * mm, 30 * mm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f3a5f")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#c8d0da")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f4f7fb")]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(table)

        notes = [m for hits in rec.fields.values() for m in hits if m.note]
        if notes:
            story.append(Spacer(1, 3 * mm))
            for m in notes:
                story.append(Paragraph(
                    f"Note on {escape(_latin1_safe(m.label))}: "
                    f"{escape(_latin1_safe(m.note))}", style_meta))

        if include_raw_text and rec.full_text:
            story.append(Spacer(1, 4 * mm))
            story.append(Paragraph("Full OCR text", styles["Heading4"]))
            for line in _latin1_safe(rec.full_text).splitlines():
                if line.strip():
                    story.append(Paragraph(escape(line), style_mono))

        if idx < len(records):
            story.append(PageBreak())

    doc.build(story)
    return buffer.getvalue()


def summary_rows(records: Sequence[ExtractionRecord]) -> list[dict]:
    """One row per image -- the batch overview table in the main pane."""
    rows = []
    for rec in records:
        utr = rec.primary("utr")
        rows.append({
            "File": rec.filename,
            "UTR / Reference": utr or "-",
            "Type": classify_utr(utr) if utr else "-",
            "Amount": rec.primary("amount") or "-",
            "Date": rec.primary("date") or "-",
            "Paid To": rec.primary("payee") or "-",
            "Status": rec.primary("status") or "-",
            "OCR conf.": f"{rec.ocr_confidence:.0%}",
        })
    return rows


# =========================================================================== #
#  4. SELF TEST
# =========================================================================== #

SAMPLES = HERE / "samples"

# What the generated receipt says, and what we expect to get back out.
EXPECTED = {
    "utr": "312845967103",
    "amount": "4250.00",
    "payee": "Rajesh Kumar",
}


def _font(size: int, bold: bool = False):
    """Prefer a real Windows UI font; fall back to PIL's bitmap default."""
    from PIL import ImageFont

    candidates = ["segoeuib.ttf", "segoeui.ttf"] if bold else ["segoeui.ttf", "arial.ttf"]
    for name in candidates:
        try:
            return ImageFont.truetype(f"C:/Windows/Fonts/{name}", size)
        except OSError:
            continue
    return ImageFont.load_default(size)


def make_sample_receipt(path: Path) -> Path:
    """Draw a UPI-style payment confirmation screenshot."""
    from PIL import ImageDraw

    width, height = 720, 1120
    img = Image.new("RGB", (width, height), "#ffffff")
    draw = ImageDraw.Draw(img)

    draw.rectangle([0, 0, width, 240], fill="#0e7a4f")
    draw.text((60, 78), "Payment Successful", font=_font(38, bold=True), fill="#ffffff")
    draw.text((60, 140), "\u20b9 4,250.00", font=_font(46, bold=True), fill="#ffffff")

    rows = [
        ("Paid To", "Rajesh Kumar"),
        ("UPI ID", "rajesh.kumar@okhdfcbank"),
        ("From", "XXXXXX4821"),
        ("IFSC Code", "HDFC0001234"),
        ("UTR No.", "312845967103"),
        ("Transaction ID", "T2408181234567890"),
        ("Date", "18 Aug 2026"),
        ("Time", "10:42 AM"),
        ("Payment Mode", "UPI"),
        ("Status", "Success"),
    ]

    y = 320
    label_font, value_font = _font(24), _font(26, bold=True)
    for label, value in rows:
        draw.text((60, y), label, font=label_font, fill="#6b7684")
        draw.text((330, y - 2), value, font=value_font, fill="#111820")
        y += 34
        draw.line([(60, y + 18), (width - 60, y + 18)], fill="#e4e8ee", width=1)
        y += 44

    draw.text((60, y + 30), "Amount debited from your account",
              font=_font(20), fill="#8a94a6")

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)
    return path


def run_selftest() -> int:
    """Generate a receipt, run the full pipeline, assert the values come back."""
    # Windows consoles default to cp1252 and OCR output can hold any script.
    if hasattr(sys.stdout, "reconfigure"):
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    engines = available_engines()
    print(f"OCR engines available : {engines or 'NONE'}")
    if not engines:
        print("\nFAIL: install an engine first -> pip install rapidocr-onnxruntime")
        return 1

    sample = make_sample_receipt(SAMPLES / "sample_receipt.png")
    print(f"Sample receipt written : {sample}")

    print("\nRunning OCR (first run loads the model, so it is the slow one)...")
    ocr = read_image(sample, engine="auto", preprocess_mode="auto")
    print(f"Engine used            : {ocr.engine}")
    print(f"Lines read             : {len(ocr.lines)}")
    print(f"Mean OCR confidence    : {ocr.mean_conf:.1%}")

    print("\n--- OCR text ---")
    print(ocr.text)

    fields = extract_fields(
        ocr,
        enabled=tuple(DEFAULT_ENABLED) + ("upi_id", "ifsc", "mode", "time", "account"),
    )
    print("\n--- Extracted fields ---")
    for hits in fields.values():
        for rank, m in enumerate(hits):
            tag = "" if rank == 0 else f"  (alt {rank + 1})"
            print(f"  {m.label:<24}: {m.value:<28} {m.confidence:.0%} via {m.method}{tag}")
            if m.note:
                print(f"  {'':<24}  ! {m.note}")

    if fields.get("utr"):
        print(f"\nUTR classified as      : {classify_utr(fields['utr'][0].value)}")

    record = ExtractionRecord(
        filename=sample.name, fields=fields, full_text=ocr.text,
        engine=ocr.engine, ocr_confidence=ocr.mean_conf, elapsed=ocr.elapsed,
    )

    print("\n--- Export formats ---")
    outputs = {
        "docx": build_docx([record], include_raw_text=True),
        "pdf": build_pdf([record], include_raw_text=True),
        "csv": build_csv([record]),
        "json": build_json([record]),
        "txt": build_txt([record], include_raw_text=True),
    }
    for ext, blob in outputs.items():
        out_path = SAMPLES / f"sample_report.{ext}"
        out_path.write_bytes(blob)
        print(f"  {ext:<5} {len(blob):>8,} bytes  -> {out_path.name}")

    print("\n--- Accuracy check ---")
    failures = []
    for key, expected in EXPECTED.items():
        got = record.primary(key)
        ok = got.upper() == expected.upper()
        print(f"  {key:<8} expected {expected!r:<20} got {got!r:<20} "
              f"{'OK' if ok else 'MISMATCH'}")
        if not ok:
            failures.append(key)

    if failures:
        print(f"\nFAIL: {len(failures)} field(s) did not match: {', '.join(failures)}")
        return 1

    print("\nPASS: OCR, extraction and all five export formats work.")
    return 0




if __name__ == "__main__":
    sys.exit(run_selftest())
